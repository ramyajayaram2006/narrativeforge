import io
import os
import re
import json
import time
import string
import html as _html
import streamlit as st
import llm
from collections import Counter
from styles import workspace_style
from relationship_map import render_relationship_map
from database import (save_story, add_character, load_characters, delete_character,
                      update_character, add_scene, load_scenes, delete_scene, update_scene,
                      add_note, load_notes, update_note, delete_note, NOTE_CATEGORIES,
                      save_snapshot, load_snapshots, restore_snapshot, delete_snapshot,
                      add_chapter, load_chapters, update_chapter, delete_chapter,
                      assign_scene_to_chapter, save_chapter_summary)
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Premium feature modules ────────────────────────────────────────────────────
try:
    from reading_mode import show_premium_reading_mode
    _HAS_READING_MODE = True
except ImportError:
    _HAS_READING_MODE = False

try:
    from book_features import show_book_mode, show_message_editor, export_book_pdf
    _HAS_BOOK_FEATURES = True
except ImportError:
    _HAS_BOOK_FEATURES = False

try:
    from version_timeline import show_version_timeline
    _HAS_VERSION_TIMELINE = True
except ImportError:
    _HAS_VERSION_TIMELINE = False

try:
    from character_suite import show_character_suite
    _HAS_CHAR_SUITE = True
except ImportError:
    _HAS_CHAR_SUITE = False

try:
    from world_builder import show_world_builder
    _HAS_WORLD_BUILDER = True
except ImportError:
    _HAS_WORLD_BUILDER = False

try:
    from directors_cut import show_directors_cut
    _HAS_DIRECTORS_CUT = True
except ImportError:
    _HAS_DIRECTORS_CUT = False

try:
    from interactive_story import show_interactive_story
    _HAS_INTERACTIVE = True
except ImportError:
    _HAS_INTERACTIVE = False

try:
    from scrivener_import import show_scrivener_import
    _HAS_IMPORTER = True
except ImportError:
    _HAS_IMPORTER = False

try:
    from beta_reading import show_beta_panel
    _HAS_BETA = True
except ImportError:
    _HAS_BETA = False

try:
    from visual_gen import show_visual_generator
    _HAS_VISUAL = True
except ImportError:
    _HAS_VISUAL = False

try:
    from plot_tools import show_plot_tools
    _HAS_PLOT_TOOLS = True
except ImportError:
    _HAS_PLOT_TOOLS = False

try:
    from screenplay import show_screenplay_tab
    _HAS_SCREENPLAY = True
except ImportError:
    _HAS_SCREENPLAY = False

try:
    from analytics import show_analytics_dashboard, record_session_words, show_full_analytics
    _HAS_ANALYTICS = True
except ImportError:
    _HAS_ANALYTICS = False

try:
    from grammar import show_grammar_checker
    _HAS_GRAMMAR = True
except ImportError:
    _HAS_GRAMMAR = False

try:
    from export import show_export_panel
    _HAS_EXPORT = True
except ImportError:
    _HAS_EXPORT = False

try:
    from achievements import (check_and_award, show_new_achievement,
                               show_achievements_panel, record_chars,
                               record_scenes, record_notes, record_analysis, record_export)
    _HAS_ACHIEVEMENTS = True
except ImportError:
    _HAS_ACHIEVEMENTS = False

try:
    from ai_tools import show_ai_tools_panel
    _HAS_AI_TOOLS = True
except ImportError:
    _HAS_AI_TOOLS = False

try:
    from daily_tracker import show_daily_tracker
    _HAS_TRACKER = True
except ImportError:
    _HAS_TRACKER = False

# ── Cached DB reads (TTL=60s) — eliminates repeated SQLite hits on tab switch
@st.cache_data(ttl=60, show_spinner=False)
def _cached_characters(username, story_id):
    return load_characters(username, story_id)

@st.cache_data(ttl=60, show_spinner=False)
def _cached_scenes(username, story_id):
    return load_scenes(username, story_id)

@st.cache_data(ttl=60, show_spinner=False)
def _cached_chapters(username, story_id):
    return load_chapters(username, story_id)

@st.cache_data(ttl=60, show_spinner=False)
def _cached_notes(username, story_id):
    return load_notes(username, story_id)

@st.cache_data(ttl=60, show_spinner=False)
def _cached_snapshots(username, story_id):
    return load_snapshots(username, story_id)

def _bust_cache(username=None, story_id=None):
    """Invalidate all read caches after any write."""
    _cached_characters.clear()
    _cached_scenes.clear()
    _cached_chapters.clear()
    _cached_notes.clear()
    _cached_snapshots.clear()

LLM_BACKEND  = "ollama"
OLLAMA_MODEL = (os.environ.get("NARRATIVEFORGE_MODEL")
                or os.environ.get("NARRATIVEFORGE_MODEL")   # legacy fallback
                or "llama3.2")
OLLAMA_HOST  = os.environ.get("OLLAMA_HOST", "http://localhost:11434")

PLOT_STAGES = [
    ("beginning",      "🌱 Beginning"),
    ("rising_action",  "📈 Rising Action"),
    ("climax",         "⚡ Climax"),
    ("falling_action", "📉 Falling Action"),
    ("resolution",     "🌅 Resolution"),
]

_STOPWORDS = {
    "the","a","an","and","or","but","in","on","at","to","for","of","with","by","from",
    "as","is","was","are","were","be","been","being","have","has","had","do","does","did",
    "will","would","could","should","may","might","shall","can","i","you","he","she","it",
    "we","they","me","him","her","us","them","my","your","his","its","our","their",
    "this","that","these","those","which","who","what","when","where","how","if","then",
    "so","not","no","up","out","about","into","than","more","just","also","said","like",
    "there","here","all","some","one","two","back","into","very","over","her","his"
}

# ── Sentiment vocabulary for Emotional Arc ────────────────────────────────────
_POS_WORDS = {
    "love","hope","joy","light","smile","laugh","warm","bright","free","peace",
    "brave","strong","alive","wonder","beautiful","safe","gentle","kind","win",
    "happy","triumph","glory","dawn","rise","good","soft","dream","grow","heal",
    "courage","golden","trust","laugh","celebrate","sparkle","bloom","serene",
}
_NEG_WORDS = {
    "dark","death","fear","pain","lost","curse","shadow","blood","fall","broken",
    "cold","hollow","dread","doom","die","hate","scream","trap","evil","void",
    "despair","rage","cry","wept","shatter","betrayal","grief","wound","sink",
    "cruel","bitter","terror","bleed","ghost","silent","alone","cage","burn","ash",
}

# ── Cliché / trope detector ────────────────────────────────────────────────────
_CLICHES = [
    "it was all a dream","chosen one","dark and stormy night","heart pounding",
    "eyes like stars","blood ran cold","time stood still","spine tingled",
    "against all odds","in the blink of an eye","last but not least","rule of thumb",
    "dead as a doornail","bite the bullet","cold as ice","heart of gold",
    "only time will tell","twist of fate","world turned upside down",
    "skeletons in the closet","moment of truth","too good to be true",
    "larger than life","at the end of the day","tip of the iceberg",
    "he had a bad feeling","something wasn't right","felt a chill",
    "thunder rumbled in the distance","the silence was deafening",
]


# ── Core helpers ───────────────────────────────────────────────────────────────
# Keys that are per-story but stored globally — cleared on story switch
_PER_STORY_KEYS = [
    "coach_result", "conflict_result", "plot_holes_result", "voice_check_result",
    "revision_results", "title_results", "firstline_results", "prompt_results",
    "rewrite_results", "char_suggest", "show_revision_mode",
]

def _get_story():
    story_id = st.session_state.current_story
    # Clear stale per-story AI results when story changes
    last_id = st.session_state.get("_last_active_story")
    if last_id and last_id != story_id:
        for k in _PER_STORY_KEYS:
            st.session_state.pop(k, None)
    st.session_state["_last_active_story"] = story_id
    for s in st.session_state.stories:
        if s["id"] == story_id:
            return s
    return None

def _word_count(story):
    return sum(len(m["content"].split()) for m in story.get("messages", []))

def _reading_time(story):
    """Estimated reading time in minutes at 200 wpm."""
    wc = _word_count(story)
    mins = max(1, round(wc / 200))
    return f"~{mins} min read"

def _all_prose(story):
    # FIX #8: only assistant messages — user messages must not skew health analytics
    return " ".join(m["content"] for m in story.get("messages", [])
                    if m["role"] == "assistant"
                    and not m["content"].startswith("◆")
                    and not m["content"].startswith("✍️"))

def _sentences(text):
    return [s.strip() for s in re.split(r'[.!?]+', text) if len(s.strip()) > 4]

def _is_incomplete(text):
    t = text.strip()

    if not t or t[-1] in ".!?": return False
    last = t.split()[-1].lower().rstrip(",:;")
    return last in {"and","but","or","with","from","to","as","while","when","the","a","an","of","in"}


# ── Story-Only Classification — 7-Instruction System ──────────────────────────
#
# Instruction 1: Classify every input as story or non-story
# Instruction 2: Story input → respond normally
# Instruction 3: Non-story input → redirect, never answer
# Instruction 4: Use one of three redirect tones (soft/firm/guiding)
# Instruction 5: Outside knowledge FOR a story = allowed (mixed requests)
# Instruction 6: Block math, coding, health, general knowledge, random facts
# Instruction 7: Never break the boundary — all non-story inputs redirected

# ── Instruction 1: Story keyword vocabulary ───────────────────────────────────
# If any of these appear in the input → classify as STORY (Instruction 5 handled here)
_STORY_KEYWORDS = {
    # Core narrative
    "plot","subplot","story","narrative","tale","fiction","novel","chapter",
    "arc","story arc","structure","act","scene","scenes","sequence","event",
    # Characters
    "character","characters","protagonist","antagonist","villain","hero",
    "heroine","narrator","mentor","sidekick","foil","backstory","motivation",
    "character arc","character development","persona","voice",
    # World & setting
    "worldbuilding","world","setting","lore","universe","realm","location",
    "atmosphere","environment","time period","era","dystopia","utopia",
    # Craft
    "dialogue","monologue","theme","themes","tone","mood","pacing","tension",
    "conflict","resolution","climax","foreshadowing","flashback","imagery",
    "metaphor","symbolism","point of view","pov","perspective","style",
    "writing style","genre","draft","revision","edit","editing","revise",
    "prose","description","exposition","show dont tell","subtext","irony",
    # Actions
    "write","writing","create","creating","develop","build","craft","describe",
    "rewrite","improve","strengthen","deepen","outline","plan","brainstorm",
    # Story elements
    "opening","beginning","ending","twist","reveal","stakes","motivation",
    "emotion","fear","love","betrayal","redemption","journey","quest",
    "timeline","flashforward","prologue","epilogue","chapter",
    # Genres & forms
    "fantasy","sci-fi","romance","mystery","horror","thriller","adventure",
    "historical fiction","literary fiction","short story","screenplay",
}

# ── Instruction 6: Explicitly blocked request types ───────────────────────────
_BLOCKED_PREFIXES = [
    # Math only
    "calculate ","solve for","how much is","how many is",
    "add ","subtract ","multiply ","divide ","compute ",
    # Coding only
    "write a program","write code","debug my","fix my code",
    "how to code","how do i code",
    # Narrow factual queries (not story-related)
    "what is the weather","what is the price","what is the stock",
    "translate this to","convert usd","install python","download the",
    # Health advice
    "diagnose me","what medicine","what symptom","how many calories",
]

# Question starters that signal non-story intent (when no story keyword present)
_Q_STARTERS = {
    # Only truly off-topic question starters (no story context)
    "calculate","solve","lookup","search",
}

# ── Redirect message variants (Instruction 4) ─────────────────────────────────
import random as _random

_REDIRECTS_SOFT = (
    "I'm designed to support story creation. Share a story idea and I'll help you build it. ✍️",
    "I'm here to help with storytelling and creative writing. What are you creating — a plot, a character, or a scene?",
)
_REDIRECTS_FIRM = (
    "I focus only on storytelling. Let's return to your story — what would you like to work on?",
    "I'm here to help with stories only. Tell me about your plot, character, or world.",
)
_REDIRECTS_GUIDING = (
    "I can't help with that, but I can help shape a story. What would you like to create? **Plot · Character · Scene · Dialogue · Worldbuilding**",
    "That's outside my storytelling focus. Let's keep things creative — do you have a character, scene, or plot idea to explore? ✍️",
)

def _get_redirect():
    pool = _REDIRECTS_SOFT + _REDIRECTS_FIRM + _REDIRECTS_GUIDING
    return _random.choice(pool)


# ── Instruction 1 classifier ──────────────────────────────────────────────────
def _classify_input(text):
    """
    Classifies input as 'story' or 'redirect' using the 7-instruction system.

    Instruction 5 — Mixed requests:
        Outside knowledge in story context is allowed.
        e.g. "explain fear psychology for my horror story"
        → contains 'story' and 'horror' → classified as STORY.

    Instruction 6 — Blocked types:
        Math, coding, general knowledge, health, random facts → redirect.
    """
    t = text.strip()
    if not t:
        return 'story'

    t_lower = t.lower()
    words = set(t_lower.replace(',','').replace('.','').replace('?','').split())

    # Instruction 5 — story keyword anywhere = story (even in questions)
    # e.g. "explain fear for my horror story" → has 'horror', 'story' → STORY
    if words & _STORY_KEYWORDS:
        return 'story'

    # Instruction 6 — blocked request types
    for prefix in _BLOCKED_PREFIXES:
        if t_lower.startswith(prefix):
            return 'redirect'

    # General question with no story context
    if t.endswith("?"):
        return 'redirect'
    if t.split()[0].lower().rstrip(",:;") in _Q_STARTERS:
        return 'redirect'

    # Instruction 7 — default: short prose → story (benefit of the doubt)
    return 'story'

def _is_non_story(text):
    return _classify_input(text) == 'redirect'


# ── Prompt builder ─────────────────────────────────────────────────────────────
def _build_prompt(user_input, story, mode, characters=None, scenes=None):
    genre = story["genre"]
    tone  = story["tone"].lower()
    writing_style = story.get("writing_style","").strip()
    history = story.get("messages", [])
    context = "\n".join(f"{m['role'].capitalize()}: {m['content']}"
                        for m in history[-10:]
                        if not m["content"].startswith("◆"))

    style_block = f"\nWriting voice: {writing_style}\n" if writing_style else ""

    char_block = ""
    if characters:
        lines = []
        for c in characters:
            line = f"  - {c['name']} ({c['role']}): {c['description']}"
            if c.get("speaking_style"):
                line += f" Their dialogue style: {c['speaking_style']}."
            lines.append(line)
        char_block = "\nCharacters in this story:\n" + "\n".join(lines) + "\n"

    scene_block = ""
    if scenes:
        sc = scenes[-1]
        scene_block = (f"\nCurrent scene: '{sc['title']}' — "
                       f"Location: {sc['location']} — Purpose: {sc['purpose']}\n")

    # Build clean story context — NO "User:" / "Assistant:" prefixes.
    # Those labels confuse the model into echoing them back as story text.
    history_msgs = history[-10:]
    story_so_far = ""
    for m in history_msgs:
        content = m["content"].strip()
        if content and not content.startswith("◆"):
            story_so_far += content + "\n\n"
    story_so_far = story_so_far.strip()

    system = (
        f"You are a professional fiction author writing a {genre} story with a {tone} tone.\n"
        f"RULES — follow all of them:\n"
        f"1. Write ONLY story prose. Never write 'User:', 'Assistant:', or any labels.\n"
        f"2. Never repeat or quote the writer's input — continue PAST it.\n"
        f"3. Never break the fourth wall or address the reader.\n"
        f"4. Never say 'I'll stop here' or 'Let me know' or any meta-commentary.\n"
        f"5. Match the tone: {tone}.\n"
        f"{style_block}{char_block}{scene_block}"
    )

    context_block = f"\n--- Story so far ---\n{story_so_far}\n--- End ---\n\n" if story_so_far else ""

    base = system + context_block

    if mode == "continue":
        return (base +
                f"The writer's last line: \"{user_input}\"\n\n"
                f"Continue the story naturally from where it left off. "
                f"Write 1-3 sentences of pure story prose. "
                f"Do NOT repeat the last line — write what happens NEXT.")

    elif mode == "paragraph":
        return (base +
                f"The writer's last line: \"{user_input}\"\n\n"
                f"Write the next full paragraph (3-5 sentences) advancing the story. "
                f"Pure prose only — no labels, no meta-text.")

    elif mode == "dialogue":
        char_list = ", ".join(c["name"] for c in (characters or []))
        return (base +
                f"The writer's last line: \"{user_input}\"\n\n"
                f"Write a spoken dialogue exchange (6-10 lines) between the story's characters "
                f"({char_list or 'the main characters'}). "
                f"Format each line as: CharacterName: \"What they say\"\n"
                f"Use each character's speaking style. No narration labels.")

    elif mode == "twist":
        return (base +
                f"The writer's last line: \"{user_input}\"\n\n"
                f"Write a 3-4 sentence paragraph that introduces a surprising, dramatic plot "
                f"twist or revelation that recontextualises what came before. "
                f"Make it feel earned. Pure prose only.")

    elif mode == "monologue":
        protagonist = next((c for c in (characters or []) if c.get("role") == "protagonist"), None)
        pname = protagonist["name"] if protagonist else "the protagonist"
        return (base +
                f"The writer's last line: \"{user_input}\"\n\n"
                f"Write a 3-5 sentence internal monologue from {pname}'s perspective — "
                f"their private thoughts, fears, or realisations. "
                f"Use first or close-third person. Lyrical and introspective. Pure prose only.")

    elif mode == "auto":
        resolved = "continue" if _is_incomplete(user_input) else "paragraph"
        return _build_prompt(user_input, story, resolved, characters, scenes)

    else:
        raise ValueError(f"Unknown mode: {mode!r}")

def _char_suggest_prompt(story, characters):
    excerpt = " ".join(m["content"] for m in story.get("messages",[])[-4:])[:600]
    existing_names = [c["name"] for c in characters] if characters else []
    existing_str = ", ".join(existing_names) if existing_names else "none yet"
    avoid = (f"\nIMPORTANT: Do NOT suggest any of these existing characters: {existing_str}. "
             f"The new character must have a completely different name and role.")  if existing_names else ""
    return (
        f"You are helping a writer build a character for their {story['genre']} story "
        f"(tone: {story['tone'].lower()}).\n"
        f"Story excerpt: {excerpt}\nExisting characters: {existing_str}{avoid}\n\n"
        f"Suggest ONE brand new character not already in the story. "
        f"Respond in EXACTLY this format, nothing else:\n"
        f"Name: [character name]\n"
        f"Role: [protagonist/antagonist/supporting/mentor/narrator]\n"
        f"Description: [2 sentences about personality and appearance]\n"
        f"Speaking style: [1 sentence about how they talk — pace, vocabulary, verbal habits]"
    )


def _strip_html(text):
    """Remove any HTML tags an LLM might accidentally generate."""
    return re.sub(r'<[^>]+>', '', text)


# ── AI backends ────────────────────────────────────────────────────────────────
def _stream_ollama(prompt):
    """Generator — yields text tokens for real-time streaming"""
    import requests
    try:
        with requests.post(
            f"{OLLAMA_HOST}/api/generate",
            json={"model": OLLAMA_MODEL, "prompt": prompt, "stream": True,
                  "options": {"temperature": 0.85, "num_predict": 280}},
            stream=True, timeout=90
        ) as r:
            if r.status_code == 404:
                yield f"[Model not found — run: ollama pull {OLLAMA_MODEL}]"
                return
            for line in r.iter_lines():
                if line:
                    try:
                        chunk = json.loads(line)
                        token = chunk.get("response", "")
                        if token:
                            yield _strip_html(token)
                        if chunk.get("done"): break
                    except Exception:
                        continue
    except Exception as e:
        if "Connection refused" in str(e) or "refused" in str(e).lower():
            yield "⚠️ Ollama not running. Open a terminal and run: ollama serve"
        else:
            yield f"[Error: {e}]"

llm.AI_UNAVAILABLE = "_llm.AI_UNAVAILABLE__"

def _call_full(prompt, max_tokens=120):
    """Non-streaming — for redirects, suggestions, coach, style mirror.
    Returns llm.AI_UNAVAILABLE sentinel string on failure so callers can show an error."""
    import requests
    try:
        r = requests.post(f"{OLLAMA_HOST}/api/generate",
                          json={"model": OLLAMA_MODEL, "prompt": prompt, "stream": False,
                                "options": {"temperature": 0.78, "num_predict": max_tokens}},
                          timeout=60)
        if r.status_code == 200:
            result = _strip_html(r.json().get("response", "").strip())
            return result if result else llm.AI_UNAVAILABLE
        return llm.AI_UNAVAILABLE
    except Exception:
        return llm.AI_UNAVAILABLE

def _ai_error_msg():
    """Render a standard AI-unavailable error."""
    st.error("⚠️ AI unavailable — is Ollama running? Try: `ollama serve`")


# ── Story Health ───────────────────────────────────────────────────────────────
def _reading_level_str(text):
    """String version for health panel display. Uses the full _reading_level() below."""
    # Defined after the full version — forward reference resolved at call time.
    result = _reading_level(text)
    if result is None:
        return "—"
    return f"{result['label']} (Grade {result['grade']})"

def _story_health(story):
    prose = _all_prose(story)
    words = prose.split()
    sents = _sentences(prose)
    total = len(words)
    clean = [w.lower().strip(string.punctuation) for w in words]
    freq  = Counter(w for w in clean if w and w not in _STOPWORDS and len(w) > 2)
    return {
        "total_words":    total,
        "total_sents":    len(sents),
        "avg_sent_len":   round(total / len(sents), 1) if sents else 0,
        "reading_label":  _reading_level_str(prose),
        "top_words":      freq.most_common(8),
        "word_goal":      story.get("word_goal", 0),
        "messages":       len(story.get("messages",[])),
    }

def _consistency_check(story, characters):
    prose  = _all_prose(story)
    words  = prose.split()
    issues = []
    total  = len(words) or 1

    clean = [w.lower().strip(string.punctuation) for w in words]
    freq  = Counter(w for w in clean if w and w not in _STOPWORDS and len(w) > 3)
    for word, count in freq.most_common(15):
        if count/total*100 > 3.0 and count >= 5:
            issues.append({"type":"overused","text":f'"{word}"',
                           "detail":f"used {count}x ({count/total*100:.1f}% of words)"})

    trigrams = [" ".join(words[i:i+3]).lower() for i in range(len(words)-2)]
    for phrase, count in Counter(trigrams).most_common(5):
        if count >= 3 and not any(sw in phrase.split()[0] for sw in list(_STOPWORDS)[:8]):
            issues.append({"type":"phrase","text":f'"{phrase}"',"detail":f"repeated {count}x"})

    long_s = [s for s in _sentences(prose) if len(s.split()) > 35]
    if long_s:
        issues.append({"type":"sentence","text":"Long sentences",
                       "detail":f"{len(long_s)} sentence(s) over 35 words"})

    if characters and len(characters) > 1:
        char_counts = {c["name"]: prose.lower().count(c["name"].lower()) for c in characters}
        total_m = sum(char_counts.values()) or 1
        for name, count in char_counts.items():
            if count/total_m*100 < 5 and total_m > 20:
                issues.append({"type":"character","text":name,
                               "detail":f"only mentioned {count}x — may feel absent"})
    return issues[:8]


# ── Emotional Arc ──────────────────────────────────────────────────────────────
def _emotional_arc(story):
    """
    Scores each AI message block for sentiment using positive/negative word sets.
    Returns list of dicts: {label, score (-1..1), pos, neg, word_count}
    """
    ai_msgs = [m for m in story.get("messages", [])
               if m["role"] == "assistant" and not m["content"].startswith("◆")]
    if not ai_msgs:
        return []
    results = []
    for i, msg in enumerate(ai_msgs):
        words = re.sub(r'[^\w\s]', '', msg["content"].lower()).split()
        pos = sum(1 for w in words if w in _POS_WORDS)
        neg = sum(1 for w in words if w in _NEG_WORDS)
        total = max(1, pos + neg)
        score = (pos - neg) / total  # -1 (very dark) to +1 (very bright)
        results.append({
            "label": f"Block {i+1}",
            "score": score,
            "pos": pos,
            "neg": neg,
            "word_count": len(words),
        })
    return results


# ── Cliché Detector ────────────────────────────────────────────────────────────
def _cliche_check(prose):
    """Returns list of clichés found in the prose."""
    lower = prose.lower()
    found = []
    for c in _CLICHES:
        if c in lower:
            found.append(c)
    return found


# ── AI Writing Coach ───────────────────────────────────────────────────────────
def _coach_prompt(story, characters):
    prose  = _all_prose(story)[:1800]  # keep prompt size manageable
    char_names = [c["name"] for c in characters] if characters else []
    return (
        f"You are an expert creative writing coach. Analyse the following story excerpt "
        f"(genre: {story['genre']}, tone: {story['tone']}) and provide EXACTLY four sections:\n\n"
        f"STRENGTHS: (2-3 bullet points on what is working well)\n"
        f"WEAKNESSES: (2-3 bullet points on specific problems)\n"
        f"SUGGESTIONS: (2-3 concrete, actionable improvements the writer can apply immediately)\n"
        f"NEXT_SCENE: (one sentence describing what should happen next to maintain momentum)\n\n"
        f"Characters present: {', '.join(char_names) if char_names else 'none defined yet'}.\n\n"
        f"Story excerpt:\n{prose}\n\n"
        f"Be specific, reference the actual text. Do not be generic. Keep total response under 250 words."
    )

def _parse_coach(raw):
    """Parse structured coach response into dict of sections."""
    sections = {"STRENGTHS": "", "WEAKNESSES": "", "SUGGESTIONS": "", "NEXT_SCENE": ""}
    current  = None
    for line in raw.splitlines():
        stripped = line.strip()
        for key in sections:
            if stripped.upper().startswith(key + ":") or stripped.upper() == key:
                current = key
                tail    = stripped[len(key):].lstrip(": ").strip()
                if tail:
                    sections[key] += tail + "\n"
                break
        else:
            if current and stripped:
                sections[current] += stripped + "\n"
    return {k: v.strip() for k, v in sections.items()}


# ── Markdown export ────────────────────────────────────────────────────────────
def _export_markdown(story, characters, scenes):
    lines = [f"# {story['title']}", "",
             f"**Genre:** {story['genre']}  |  **Tone:** {story['tone']}  |  "
             f"**Words:** {_word_count(story):,}", ""]
    if story.get("writing_style"):
        lines += ["## Writing Voice", story["writing_style"], ""]
    if characters:
        lines += ["## Characters", ""]
        for c in characters:
            lines.append(f"### {c['name']}  _{c['role']}_")
            lines.append(c["description"])
            if c.get("speaking_style"):
                lines.append(f"*Speaking style: {c['speaking_style']}*")
            lines.append("")
    arc = story.get("plot_arc", {})
    completed = [label for key, label in PLOT_STAGES if arc.get(key)]
    if completed:
        lines += ["## Plot Arc", " → ".join(completed), ""]
    if scenes:
        lines += ["## Scene Overview", ""]
        for i, sc in enumerate(scenes, 1):
            lines.append(f"### Scene {i}: {sc['title']}")
            if sc["location"]: lines.append(f"📍 {sc['location']}")
            if sc["purpose"]:  lines.append(f"🎯 {sc['purpose']}")
            if sc["characters"]: lines.append(f"👥 {', '.join(sc['characters'])}")
            lines.append("")
    lines += ["## Story", "", "---", ""]
    prose = " ".join(m["content"] for m in story.get("messages", [])
                     if m["role"] == "assistant" and not m["content"].startswith("◆"))
    lines.append(prose.strip() if prose else "_No content yet._")
    lines += ["", "---", "", "*Generated by NarrativeForge*"]
    return "\n".join(lines).encode("utf-8")


# ── Word export ────────────────────────────────────────────────────────────────
def _export_docx(story, characters, scenes, scene_chapters=False):
    """
    Export story to .docx.
    scene_chapters=True  → NEW FEATURE: each scene gets its own chapter heading,
                           with the prose sliced proportionally across scenes.
    scene_chapters=False → original flat export (all prose in one block).
    """
    doc = Document()
    t = doc.add_heading(story["title"], 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m = doc.add_paragraph()
    m.add_run(f"Genre: {story['genre']}  |  Tone: {story['tone']}  |  Words: {_word_count(story)}").italic = True
    m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    if story.get("writing_style"):
        doc.add_heading("Writing Voice", 1)
        doc.add_paragraph(story["writing_style"])
        doc.add_paragraph()

    if characters:
        doc.add_heading("Characters", 1)
        for c in characters:
            p = doc.add_paragraph()
            p.add_run(f"{c['name']}").bold = True
            p.add_run(f"  ({c['role']})")
            doc.add_paragraph(c["description"])
            if c.get("speaking_style"):
                sp = doc.add_paragraph()
                sp.add_run("Speaking style: ").bold = True
                sp.add_run(c["speaking_style"])
        doc.add_paragraph()

    arc = story.get("plot_arc", {})
    completed = [label for key, label in PLOT_STAGES if arc.get(key)]
    if completed:
        doc.add_heading("Plot Arc", 1)
        doc.add_paragraph(" → ".join(completed))
        doc.add_paragraph()

    # Collect all AI-generated prose sentences
    all_msgs = [m for m in story.get("messages", [])
                if m["role"] == "assistant" and not m["content"].startswith("◆")]
    prose_full = " ".join(m["content"] for m in all_msgs).strip()

    if scene_chapters and scenes and prose_full:
        # ── NEW FEATURE: Scene-structured chapters ─────────────────────────
        # Distribute prose sentences evenly across scenes
        import re as _re
        sentences = [s.strip() for s in _re.split(r'(?<=[.!?])\s+', prose_full) if s.strip()]
        n_scenes  = len(scenes)
        chunk     = max(1, len(sentences) // n_scenes)

        doc.add_heading("Story", 1)
        for idx, sc in enumerate(scenes):
            # Chapter heading: Scene title as H2
            doc.add_heading(f"Chapter {idx + 1}: {sc['title']}", 2)
            # Scene meta line
            meta_parts = []
            if sc.get("location"): meta_parts.append(f"📍 {sc['location']}")
            if sc.get("purpose"):  meta_parts.append(f"🎯 {sc['purpose']}")
            if sc.get("characters"): meta_parts.append(f"👥 {', '.join(sc['characters'])}")
            if meta_parts:
                mp = doc.add_paragraph(" · ".join(meta_parts))
                mp.runs[0].italic = True
            doc.add_paragraph()

            # Slice of prose for this chapter
            start = idx * chunk
            end   = start + chunk if idx < n_scenes - 1 else len(sentences)
            slice_text = " ".join(sentences[start:end])
            doc.add_paragraph(slice_text if slice_text else "(No content yet for this chapter.)")
            doc.add_paragraph()
    else:
        # ── Original flat export ───────────────────────────────────────────
        if scenes:
            doc.add_heading("Scene Overview", 1)
            for i, sc in enumerate(scenes, 1):
                p = doc.add_paragraph()
                p.add_run(f"Scene {i}: {sc['title']}").bold = True
                if sc["location"]: doc.add_paragraph(f"Location: {sc['location']}")
                if sc["purpose"]:  doc.add_paragraph(f"Purpose: {sc['purpose']}")
                if sc["characters"]: doc.add_paragraph(f"Characters: {', '.join(sc['characters'])}")
            doc.add_paragraph()

        doc.add_heading("Story", 1)
        doc.add_paragraph("─" * 50)
        doc.add_paragraph(prose_full if prose_full else "No content yet.")
        doc.add_paragraph()

    f = doc.add_paragraph()
    f.add_run("Generated by NarrativeForge").italic = True
    f.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── Sidebar panels ─────────────────────────────────────────────────────────────
def _sidebar_settings(story):
    wc   = _word_count(story)
    goal = story.get("word_goal", 0)
    pct  = min(100, int(wc / goal * 100)) if goal else None

    # ── Premium SVG Progress Ring ──────────────────────────────────────────────
    if pct is not None:
        r, cx, cy = 40, 54, 54
        circumference = 2 * 3.14159 * r
        dash_offset   = circumference * (1 - pct / 100)
        ring_color    = "#4ADE80" if pct >= 100 else "#4D6BFE" if pct >= 50 else "#fbbf24"
        ring_svg = f"""
        <div style='display:flex;align-items:center;gap:14px;margin-bottom:8px;'>
          <svg width="108" height="108" viewBox="0 0 108 108" xmlns="http://www.w3.org/2000/svg">
            <circle cx="{cx}" cy="{cy}" r="{r}" fill="none"
              stroke="rgba(77,107,254,0.12)" stroke-width="10"/>
            <circle cx="{cx}" cy="{cy}" r="{r}" fill="none"
              stroke="{ring_color}" stroke-width="10" stroke-linecap="round"
              stroke-dasharray="{circumference:.1f}"
              stroke-dashoffset="{dash_offset:.1f}"
              transform="rotate(-90 {cx} {cy})"/>
            <text x="{cx}" y="{cy-6}" text-anchor="middle"
              font-size="16" font-weight="800" fill="{ring_color}"
              font-family="Inter,sans-serif">{pct}%</text>
            <text x="{cx}" y="{cy+10}" text-anchor="middle"
              font-size="9" fill="#6B7080" font-family="Inter,sans-serif">of goal</text>
          </svg>
          <div>
            <div style='font-size:1.6rem;font-weight:800;color:var(--primary);line-height:1;'>{wc:,}</div>
            <div style='font-size:0.72rem;color:#6B7080;'>words written</div>
            <div style='font-size:0.72rem;color:#6B7080;margin-top:2px;'>goal: {goal:,}</div>
            <div style='font-size:0.72rem;color:#6B7080;'>{len(story.get("messages",[]))} messages</div>
          </div>
        </div>"""
        st.markdown(ring_svg, unsafe_allow_html=True)
    else:
        st.markdown(f"""
            <div class='wc-box'>
                <div class='wc-number'>{wc:,}</div>
                <div class='wc-label'>Words Written</div>
            </div>
        """, unsafe_allow_html=True)
        st.markdown(f"<div class='wc-sub'>{len(story.get('messages',[]))} messages</div>", unsafe_allow_html=True)

    st.markdown("---")

    new_title = st.text_input("Story Title", value=story["title"],
                               key="ws_title", max_chars=100)
    if new_title.strip() and new_title.strip() != story["title"]:
        story["title"] = new_title.strip()
        save_story(st.session_state.username, story)  # FIX #3: persist immediately

    genres = ["Fantasy","Sci-Fi","Mystery","Romance","Horror","Adventure","Thriller","Historical"]
    new_genre = st.selectbox("Genre", genres,
        index=genres.index(story.get("genre","Fantasy")) if story.get("genre") in genres else 0,
        key="ws_genre")
    if new_genre != story.get("genre"):
        story["genre"] = new_genre
        save_story(st.session_state.username, story)  # FIX #14: persist immediately

    # Genre craft tips
    _show_genre_tips(story.get("genre", "Fantasy"))

    tones = ["Dark","Light","Emotional","Humorous","Serious","Suspenseful","Whimsical"]
    new_tone = st.selectbox("Tone", tones,
        index=tones.index(story.get("tone","Light")) if story.get("tone") in tones else 1,
        key="ws_tone")
    if new_tone != story.get("tone"):
        story["tone"] = new_tone
        save_story(st.session_state.username, story)  # FIX #14: persist immediately

    st.markdown("<div style='margin-top:10px;font-size:0.78rem;color:#6B7080;font-weight:600;text-transform:uppercase;letter-spacing:0.08em;'>✍️ Writing Voice</div>", unsafe_allow_html=True)
    new_style = st.text_area("_", value=story.get("writing_style",""),
                              placeholder="e.g. lyrical, slow-burn, first-person past tense, rich sensory details",
                              height=75, key="ws_style", label_visibility="collapsed")
    if new_style != story.get("writing_style",""):
        story["writing_style"] = new_style
        save_story(st.session_state.username, story)

    goal_input = st.number_input("🎯 Word Goal", min_value=0, max_value=200000,
                                  value=int(story.get("word_goal",0)), step=500, key="ws_goal")
    if goal_input != story.get("word_goal",0):
        story["word_goal"] = goal_input
        save_story(st.session_state.username, story)

    st.markdown(f"<div style='color:#6B7080;font-size:0.75rem;margin-top:6px;'>🤖 AI: <code style='color:#4D6BFE;'>{LLM_BACKEND}</code></div>", unsafe_allow_html=True)

    # ── Style Mirror ───────────────────────────────────────────────────────────
    st.markdown("---")
    st.markdown("<div style='font-size:0.72rem;color:#6B7080;text-transform:uppercase;"
                "letter-spacing:0.08em;font-weight:600;margin-bottom:6px;'>🪞 Style Mirror</div>",
                unsafe_allow_html=True)
    current_dna = story.get("style_dna", "")
    if current_dna:
        st.markdown(
            f"<div style='background:rgba(77,107,254,0.06);border:1px solid rgba(77,107,254,0.18);"
            f"border-radius:8px;padding:8px 12px;font-size:0.78rem;color:#A8BCFF;"
            f"font-style:italic;line-height:1.5;margin-bottom:8px;'>"
            f"🪞 {_html.escape(current_dna[:140])}{'…' if len(current_dna)>140 else ''}</div>",
            unsafe_allow_html=True)
    with st.expander("✨ Extract from author sample"):
        sample = st.text_area("Paste 2-3 paragraphs from your favourite author",
                              height=110, key="style_sample_input",
                              placeholder="Paste prose here — NarrativeForge will extract the style signature and inject it into every AI generation.")
        if st.button("🪞 Extract Style DNA", use_container_width=True, type="primary", key="extract_style_btn"):
            if len(sample.strip()) < 80:
                st.warning("Paste at least 2 paragraphs for a useful style analysis.")
            else:
                with st.spinner("Analysing prose style…"):
                    dna_prompt = (
                        "Analyse the writing style of the following prose excerpt. "
                        "Describe ONLY the style — sentence length, rhythm, vocabulary level, "
                        "POV, tense, tone, figurative language, and distinctive quirks. "
                        "Write ONE concise paragraph (40-60 words) in the imperative: "
                        "'Write in [style details]…'. Do not summarise the content. "
                        "Be specific and technical.\n\n"
                        f"Excerpt:\n{sample[:1200]}"
                    )
                    dna = _call_full(dna_prompt, max_tokens=120)
                if dna:
                    story["style_dna"]       = dna
                    story["writing_style"]   = dna
                    save_story(st.session_state.username, story)
                    st.success("Style extracted and applied! Every generation will now mirror this voice.")
                    _bust_cache()
                    st.rerun()
                else:
                    st.error("Could not extract style — is Ollama running?")
        if current_dna and st.button("🗑️ Clear Style DNA", use_container_width=True, key="clear_dna_btn"):
            story["style_dna"]     = ""
            story["writing_style"] = ""
            save_story(st.session_state.username, story)
            _bust_cache()
            st.rerun()

    st.markdown("---")

    chars  = _cached_characters(st.session_state.username, story["id"])
    scenes = _cached_scenes(st.session_state.username, story["id"])
    if story.get("messages"):
        safe_title = story['title'].replace(' ', '_')
        st.markdown(f"<div style='font-size:0.72rem;color:#6B7080;margin-bottom:6px;text-align:center;'>"
                    f"⏱ {_reading_time(story)}</div>", unsafe_allow_html=True)
        d1, d2 = st.columns(2)
        with d1:
            st.download_button("📥 .docx",
                data=_export_docx(story, chars, scenes, scene_chapters=False),
                file_name=f"{safe_title}_NarrativeForge.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True)
        with d2:
            st.download_button("📝 .md",
                data=_export_markdown(story, chars, scenes),
                file_name=f"{safe_title}_NarrativeForge.md",
                mime="text/markdown",
                use_container_width=True)
        if scenes:
            st.download_button("📖 Chapters .docx",
                data=_export_docx(story, chars, scenes, scene_chapters=True),
                file_name=f"{safe_title}_Chapters_NarrativeForge.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                help="Each scene becomes a titled chapter")
        # PDF export — requires reportlab
        try:
            pdf_data = _export_pdf(story, chars, scenes)
            st.download_button("PDF",
                data=pdf_data,
                file_name=f"{safe_title}_NarrativeForge.pdf",
                mime="application/pdf",
                use_container_width=True,
                help="Formatted PDF export")
        except ImportError:
            st.markdown(
                "<div style='font-size:0.68rem;color:#6B7080;padding:4px 0;'>"
                "PDF: run <code>pip install reportlab</code></div>",
                unsafe_allow_html=True)

    if st.button("🗑️ Clear Chat", use_container_width=True):
        story["messages"] = []
        save_story(st.session_state.username, story)
        st.session_state.show_cowrite_options = False
        _bust_cache()
        st.rerun()

    st.markdown("---")
    # ── Pomodoro Timer ────────────────────────────────────────────────────
    _pomodoro_widget(story["id"])

    st.markdown("---")
    # ── First Draft Mode ──────────────────────────────────────────────────
    fd_key = f"_first_draft_{story['id']}"
    is_fd  = st.session_state.get(fd_key, False)
    fd_label = "🔒 First Draft ON — no deleting!" if is_fd else "✍️ First Draft Mode"
    if st.toggle(fd_label, value=is_fd, key=f"fd_toggle_{story['id']}"):
        st.session_state[fd_key] = True
    else:
        st.session_state[fd_key] = False
    if is_fd:
        st.markdown(
            "<div style='font-size:0.70rem;color:var(--warning);padding:4px 0;'>"
            "Backspace &amp; Delete are disabled. Keep writing forward!</div>",
            unsafe_allow_html=True)
        _inject_first_draft_mode()

    st.markdown("---")
    from auth import show_account_settings, _full_logout
    show_account_settings()
    st.markdown("---")
    from styles import theme_toggle_widget
    theme_toggle_widget()
    st.markdown("---")
    if st.button("Sign Out", use_container_width=True):
        _full_logout()
        st.rerun()


def _sidebar_characters(story):
    username   = st.session_state.username
    # FIX: Always load fresh from DB — never use cached list
    characters = load_characters(username, story["id"])
    roles_list = ["protagonist","antagonist","supporting","mentor","narrator"]

    if characters:
        for c in characters:
            voice_html = f"<div class='char-voice'>🗣 {_html.escape(c['speaking_style'])}</div>" if c.get("speaking_style") else ""
            arc_html   = (f"<div style='font-size:0.70rem;color:var(--primary);font-style:italic;"
                          f"margin-top:3px;'>📈 {_html.escape(c['arc_notes'][:80])}{'…' if len(c.get('arc_notes',''))>80 else ''}</div>"
                          if c.get("arc_notes") else "")
            st.markdown(f"""
                <div class='char-card'>
                    <div class='char-name'>{_html.escape(c['name'])}</div>
                    <div class='char-role'>{_html.escape(c['role'])}</div>
                    <div class='char-desc'>{_html.escape(c['description'])}</div>
                    {voice_html}{arc_html}
                </div>
            """, unsafe_allow_html=True)

            # Edit + Delete buttons side by side
            col_e, col_d = st.columns(2)
            with col_e:
                if st.button("✏️ Edit", key=f"editchar_{c['id']}", use_container_width=True):
                    # Toggle: click again to close
                    if st.session_state.get("editing_char") == c["id"]:
                        st.session_state.pop("editing_char", None)
                    else:
                        st.session_state["editing_char"] = c["id"]
                    st.rerun()
            with col_d:
                if not st.session_state.get(f"_confirm_del_char_{c['id']}"):
                    if st.button("Delete", key=f"delchar_{c['id']}", use_container_width=True):
                        st.session_state[f"_confirm_del_char_{c['id']}"] = True
                        st.rerun()
                else:
                    st.markdown(
                        "<div style='background:rgba(229,115,115,0.08);"
                        "border:1px solid rgba(229,115,115,0.25);border-radius:6px;"
                        "padding:4px 8px;font-size:0.68rem;color:#FFCDD2;"
                        "margin-bottom:4px;'>Delete this character?</div>",
                        unsafe_allow_html=True)
                    ca, cb = st.columns(2)
                    with ca:
                        if st.button("Yes", key=f"confirm_delchar_{c['id']}", use_container_width=True):
                            delete_character(c["id"], username)
                            st.session_state.pop("editing_char", None)
                            st.session_state.pop(f"_confirm_del_char_{c['id']}", None)
                            _bust_cache()
                            st.rerun()
                    with cb:
                        if st.button("No", key=f"cancel_delchar_{c['id']}", use_container_width=True):
                            st.session_state.pop(f"_confirm_del_char_{c['id']}", None)
                            _bust_cache()
                            st.rerun()

            # Inline edit form — only shown for the selected character
            if st.session_state.get("editing_char") == c["id"]:
                st.markdown("<div style='background:rgba(77,107,254,0.06);border:1px solid rgba(77,107,254,0.20);border-radius:8px;padding:10px;margin:6px 0;'>", unsafe_allow_html=True)
                e_name  = st.text_input("Name",           value=c["name"],  key=f"ename_{c['id']}")
                e_role  = st.selectbox("Role", roles_list,
                            index=roles_list.index(c["role"]) if c["role"] in roles_list else 2,
                            key=f"erole_{c['id']}")
                e_desc  = st.text_area("Description",     value=c["description"], height=70, key=f"edesc_{c['id']}")
                e_voice = st.text_input("Speaking Style",  value=c.get("speaking_style",""), key=f"evoice_{c['id']}")
                e_arc   = st.text_area("📈 Arc Notes",
                            value=c.get("arc_notes",""),
                            height=55, key=f"earc_{c['id']}",
                            placeholder="How does this character change across the story?")
                sv1, sv2 = st.columns(2)
                with sv1:
                    if st.button("💾 Save", key=f"savechar_{c['id']}", use_container_width=True):
                        update_character(c["id"], username, e_name, e_role, e_desc, e_voice, e_arc)  # FIX #1: IDOR + arc_notes
                        st.session_state.pop("editing_char", None)
                        _bust_cache()
                        st.rerun()
                with sv2:
                    if st.button("✖ Cancel", key=f"cancelchar_{c['id']}", use_container_width=True):
                        st.session_state.pop("editing_char", None)
                        _bust_cache()
                        st.rerun()
                st.markdown("</div>", unsafe_allow_html=True)
    else:
        st.markdown("<div style='color:#6B7080;font-size:0.82rem;text-align:center;padding:12px;'>No characters yet.</div>", unsafe_allow_html=True)

    st.markdown("---")

    # AI Suggest — always queries fresh list, always clears previous suggestion
    if st.button("✨ AI Suggest", use_container_width=True, type="primary"):
        for k in ["char_suggest","char_name","char_role","char_desc","char_voice"]:
            st.session_state.pop(k, None)
        fresh_chars = load_characters(username, story["id"])
        with st.spinner("Thinking of a new character..."):
            raw = _call_full(_char_suggest_prompt(story, fresh_chars), max_tokens=160)
        s = {"name":"","role":"supporting","description":"","speaking_style":""}
        for line in raw.splitlines():
            ll = line.lower()
            if ll.startswith("name:"):        s["name"]          = line.split(":",1)[1].strip()
            elif ll.startswith("role:"):
                r = line.split(":",1)[1].strip().lower()
                s["role"] = r if r in roles_list else "supporting"
            elif ll.startswith("description:"): s["description"] = line.split(":",1)[1].strip()
            elif "speaking" in ll and ":" in ll: s["speaking_style"] = line.split(":",1)[1].strip()
        st.session_state["char_suggest"] = s
        st.rerun()

    sug = st.session_state.get("char_suggest", {})
    with st.expander("➕ Add Character", expanded=bool(sug)):
        c_name  = st.text_input("Name",  value=sug.get("name",""),  key="char_name")
        c_role  = st.selectbox("Role", roles_list,
                    index=roles_list.index(sug.get("role","supporting")), key="char_role")
        c_desc  = st.text_area("Description", value=sug.get("description",""),
                    height=70, key="char_desc", placeholder="Personality, appearance, backstory...")
        c_voice = st.text_input("Speaking Style", value=sug.get("speaking_style",""),
                    key="char_voice", placeholder="e.g. formal and slow, speaks in riddles...")
        c_arc   = st.text_area("📈 Arc Notes", value="", height=55, key="char_arc",
                    placeholder="How does this character change across the story?")
        if st.button("💾 Save Character", use_container_width=True):
            if c_name.strip():
                add_character(username, story["id"], c_name, c_role, c_desc, c_voice, c_arc)
                for k in ["char_suggest","char_name","char_role","char_desc","char_voice"]:
                    st.session_state.pop(k, None)
                _bust_cache()
                st.rerun()


def _sidebar_scenes(story):
    username   = st.session_state.username
    # FIX: Always load fresh from DB
    scenes     = _cached_scenes(username, story["id"])
    characters = _cached_characters(username, story["id"])
    # FIX: Always reflect current character list — never stale
    char_names = [c["name"] for c in characters]

    if scenes:
        for sc in scenes:
            chars_html = f"<div class='scene-meta'>👥 {_html.escape(', '.join(sc['characters']))}</div>" if sc["characters"] else ""
            st.markdown(f"""
                <div class='scene-card'>
                    <div class='scene-title'>Scene {sc['order']}: {_html.escape(sc['title'])}</div>
                    <div class='scene-meta'>📍 {_html.escape(sc['location'] or '—')} · 🎯 {_html.escape(sc['purpose'] or '—')}</div>
                    {chars_html}
                </div>
            """, unsafe_allow_html=True)

            se1, se2 = st.columns(2)
            with se1:
                if st.button("✏️ Edit", key=f"editscene_{sc['id']}", use_container_width=True):
                    if st.session_state.get("editing_scene") == sc["id"]:
                        st.session_state.pop("editing_scene", None)
                    else:
                        st.session_state["editing_scene"] = sc["id"]
                    st.rerun()
            with se2:
                if not st.session_state.get(f"_confirm_del_scene_{sc['id']}"):
                    if st.button("Delete", key=f"delscene_{sc['id']}", use_container_width=True):
                        st.session_state[f"_confirm_del_scene_{sc['id']}"] = True
                        st.rerun()
                else:
                    st.markdown(
                        "<div style='background:rgba(229,115,115,0.08);"
                        "border:1px solid rgba(229,115,115,0.25);border-radius:6px;"
                        "padding:4px 8px;font-size:0.68rem;color:#FFCDD2;margin-bottom:4px;'>"
                        "Delete this scene?</div>",
                        unsafe_allow_html=True)
                    sa, sb = st.columns(2)
                    with sa:
                        if st.button("Yes", key=f"confirm_delscene_{sc['id']}", use_container_width=True):
                            delete_scene(sc["id"], username)
                            st.session_state.pop("editing_scene", None)
                            st.session_state.pop(f"_confirm_del_scene_{sc['id']}", None)
                            _bust_cache()
                            st.rerun()
                    with sb:
                        if st.button("No", key=f"cancel_delscene_{sc['id']}", use_container_width=True):
                            st.session_state.pop(f"_confirm_del_scene_{sc['id']}", None)
                            _bust_cache()
                            st.rerun()

            # Inline edit form for this scene
            if st.session_state.get("editing_scene") == sc["id"]:
                st.markdown("<div style='background:rgba(77,107,254,0.06);border:1px solid rgba(77,107,254,0.20);border-radius:8px;padding:10px;margin:6px 0;'>", unsafe_allow_html=True)
                es_title    = st.text_input("Title",    value=sc["title"],    key=f"estitle_{sc['id']}")
                es_location = st.text_input("Location", value=sc["location"], key=f"esloc_{sc['id']}")
                es_purpose  = st.text_input("Purpose",  value=sc["purpose"],  key=f"espurp_{sc['id']}")
                # FIX: Character multiselect uses CURRENT character list, pre-selects existing
                current_in_scene = [ch for ch in sc["characters"] if ch in char_names]
                es_chars = st.multiselect("Characters", char_names,
                             default=current_in_scene, key=f"eschars_{sc['id']}") if char_names else []
                ss1, ss2 = st.columns(2)
                with ss1:
                    if st.button("💾 Save", key=f"savescene_{sc['id']}", use_container_width=True):
                        update_scene(sc["id"], username, es_title, es_location, es_purpose, es_chars)  # FIX #1: IDOR
                        st.session_state.pop("editing_scene", None)
                        _bust_cache()
                        st.rerun()
                with ss2:
                    if st.button("✖ Cancel", key=f"cancelscene_{sc['id']}", use_container_width=True):
                        st.session_state.pop("editing_scene", None)
                        _bust_cache()
                        st.rerun()
                st.markdown("</div>", unsafe_allow_html=True)
    else:
        st.markdown("<div style='color:#6B7080;font-size:0.82rem;text-align:center;padding:12px;'>No scenes yet.</div>", unsafe_allow_html=True)

    st.markdown("---")
    with st.expander("➕ Add Scene"):
        s_title    = st.text_input("Title",    key="s_title")
        s_location = st.text_input("Location", key="s_loc")
        s_purpose  = st.text_input("Purpose",  key="s_purpose",
                                   placeholder="e.g. Introduce conflict")
        # FIX: Multiselect always shows current characters — empty list if none added yet
        s_chars = st.multiselect("Characters in this scene", char_names, key="s_chars") if char_names else []
        if not char_names:
            st.caption("Add characters in the Cast tab first to assign them to scenes.")
        if st.button("💾 Save Scene", use_container_width=True):
            if s_title.strip():
                add_scene(username, story["id"], s_title, s_location, s_purpose, s_chars)
                for k in ["s_title","s_loc","s_purpose","s_chars"]:
                    st.session_state.pop(k, None)
                _bust_cache()
                st.rerun()


def _sidebar_arc(story):
    if _HAS_PLOT_TOOLS:
        show_plot_tools(story, st.session_state.username)
    else:
        # Fallback: simple checkboxes
        arc     = story.get("plot_arc", {})
        changed = False
        for key, label in PLOT_STAGES:
            val = st.checkbox(label, value=arc.get(key,False), key=f"arc_{story['id']}_{key}")
            if val != arc.get(key,False): arc[key] = val; changed = True
        if changed:
            story["plot_arc"] = arc
            save_story(st.session_state.username, story)


def _sidebar_health(story, characters):
    if not story.get("messages"):
        st.markdown("<div style='color:#6B7080;font-size:0.82rem;text-align:center;padding:20px;'>Write some story first.</div>", unsafe_allow_html=True)
        return

    h      = _story_health(story)
    issues = _consistency_check(story, characters)

    # ── Stats grid ────────────────────────────────────────────────────────────
    rt = _reading_time(story)
    st.markdown(f"""
        <div class='health-grid'>
            <div class='health-stat'>
                <div class='hs-num'>{h['avg_sent_len']}</div>
                <div class='hs-lbl'>Avg Sent.<br>Length</div>
            </div>
            <div class='health-stat'>
                <div class='hs-num' style='font-size:0.85rem;'>{h['reading_label']}</div>
                <div class='hs-lbl'>Reading<br>Level</div>
            </div>
            <div class='health-stat'>
                <div class='hs-num' style='font-size:0.78rem;'>{rt}</div>
                <div class='hs-lbl'>Est. Read<br>Time</div>
            </div>
            <div class='health-stat'>
                <div class='hs-num'>{h['total_sents']}</div>
                <div class='hs-lbl'>Total<br>Sentences</div>
            </div>
        </div>
    """, unsafe_allow_html=True)
    st.markdown("---")

    # ── Emotional Arc ──────────────────────────────────────────────────────────
    arc_data = _emotional_arc(story)
    if arc_data:
        st.markdown("<div style='font-size:0.72rem;color:#6B7080;text-transform:uppercase;letter-spacing:0.08em;margin-bottom:8px;font-weight:600;'>🎭 Emotional Arc</div>", unsafe_allow_html=True)
        for pt in arc_data:
            score    = pt["score"]           # -1 (dark) to +1 (bright)
            pct      = int((score + 1) / 2 * 100)  # map to 0-100
            if score > 0.2:   color, mood = "#4D6BFE", "hopeful"
            elif score < -0.2: color, mood = "#f87171", "tense"
            else:              color, mood = "#fbbf24", "neutral"
            st.markdown(f"""
                <div style='display:flex;align-items:center;gap:6px;margin:3px 0;'>
                    <div style='width:52px;font-size:0.68rem;color:#6B7080;white-space:nowrap;'>{_html.escape(pt["label"])}</div>
                    <div style='flex:1;background:rgba(77,107,254,0.08);border-radius:3px;height:7px;'>
                        <div style='width:{pct}%;background:{color};height:7px;border-radius:3px;'></div>
                    </div>
                    <div style='width:46px;font-size:0.68rem;color:{color};text-align:right;'>{mood}</div>
                </div>
            """, unsafe_allow_html=True)
        st.markdown("---")

    # ── Word Frequency ─────────────────────────────────────────────────────────
    if h["top_words"]:
        st.markdown("<div style='font-size:0.72rem;color:#6B7080;text-transform:uppercase;letter-spacing:0.08em;margin-bottom:8px;font-weight:600;'>📊 Word Frequency</div>", unsafe_allow_html=True)
        max_count = h["top_words"][0][1] if h["top_words"] else 1
        for word, count in h["top_words"]:
            total_w = max(h["total_words"], 1)
            density = count / total_w * 100
            bar_w   = int(count / max_count * 100)
            color   = "#e11d48" if density > 3.5 else "#4D6BFE"
            st.markdown(f"""
                <div style='display:flex;align-items:center;gap:8px;margin:4px 0;'>
                    <div style='width:65px;font-size:0.78rem;color:var(--text-primary);font-weight:500;overflow:hidden;white-space:nowrap;'>{_html.escape(word)}</div>
                    <div style='flex:1;background:rgba(77,107,254,0.08);border-radius:3px;height:5px;'>
                        <div style='width:{bar_w}%;background:{color};height:5px;border-radius:3px;'></div>
                    </div>
                    <div style='width:28px;font-size:0.72rem;color:#6B7080;text-align:right;'>{count}x</div>
                </div>
            """, unsafe_allow_html=True)

    # ── Vocabulary Richness ────────────────────────────────────────────────────
    vocab = _vocab_stats(story)
    if vocab:
        st.markdown("---")
        st.markdown("<div style='font-size:0.72rem;color:#6B7080;text-transform:uppercase;"
                    "letter-spacing:0.08em;margin-bottom:8px;font-weight:600;'>📖 Vocabulary Richness</div>",
                    unsafe_allow_html=True)
        v1, v2 = st.columns(2)
        v1.metric("Unique Words", f"{vocab['unique']:,}")
        v2.metric("Richness", f"{vocab['richness']}%", help="Unique ÷ total words × 100")
        v3, v4 = st.columns(2)
        v3.metric("Dialogue", f"{vocab['dialogue_pct']}%", help="% of words inside quotes")
        v4.metric("Avg Word Len", f"{vocab['avg_word_len']} ch")
        richness = vocab["richness"]
        if richness >= 65:
            verdict = "🌟 Excellent vocabulary diversity"
            color   = "#4D6BFE"
        elif richness >= 45:
            verdict = "✅ Good — healthy word variety"
            color   = "#4ade80"
        elif richness >= 30:
            verdict = "⚠️ Moderate — some repetition"
            color   = "#fbbf24"
        else:
            verdict = "🔁 Low — many repeated words"
            color   = "#f87171"
        st.markdown(
            f"<div style='font-size:0.75rem;color:{color};padding:6px 0;'>{verdict}</div>",
            unsafe_allow_html=True)

    # ── Cliché Detector ────────────────────────────────────────────────────────
    prose   = _all_prose(story)
    cliches = _cliche_check(prose)
    st.markdown("---")
    st.markdown("<div style='font-size:0.72rem;color:#6B7080;text-transform:uppercase;letter-spacing:0.08em;margin-bottom:8px;font-weight:600;'>⚠️ Clichés Detected</div>", unsafe_allow_html=True)
    if not cliches:
        st.markdown("<div style='color:#4D6BFE;font-size:0.82rem;padding:4px;'>✅ No common clichés found!</div>", unsafe_allow_html=True)
    else:
        for c in cliches[:6]:
            st.markdown(f"""
                <div class='issue-card' style='border-left-color:#fbbf24;'>
                    <div style='font-size:0.78rem;color:#fcd34d;font-weight:600;'>⚠️ &ldquo;{_html.escape(c)}&rdquo;</div>
                    <div style='font-size:0.68rem;color:#6B7080;margin-top:2px;'>Common trope — consider a fresh take</div>
                </div>
            """, unsafe_allow_html=True)

    # ── Consistency issues ─────────────────────────────────────────────────────
    st.markdown("---")
    st.markdown("<div style='font-size:0.72rem;color:#6B7080;text-transform:uppercase;letter-spacing:0.08em;margin-bottom:8px;font-weight:600;'>🔍 Consistency</div>", unsafe_allow_html=True)
    if not issues:
        st.markdown("<div style='color:#4D6BFE;font-size:0.82rem;text-align:center;padding:8px;'>✅ No issues found!</div>", unsafe_allow_html=True)
    else:
        icons  = {"overused": "🔁", "phrase": "💬", "sentence": "📏", "character": "👤"}
        colors = {"overused": "#fca5a5", "phrase": "#fcd34d", "sentence": "#93c5fd", "character": "#c4b5fd"}
        for issue in issues:
            ic  = icons.get(issue["type"], "⚠️")
            clr = colors.get(issue["type"], "#A8BCFF")
            st.markdown(f"""
                <div class='issue-card' style='border-left-color:{clr};'>
                    <div style='font-size:0.8rem;color:var(--text-primary);font-weight:600;'>{ic} {_html.escape(issue['text'])}</div>
                    <div style='font-size:0.72rem;color:#6B7080;margin-top:2px;'>{_html.escape(issue['detail'])}</div>
                </div>
            """, unsafe_allow_html=True)

    # ── AI Writing Coach ───────────────────────────────────────────────────────
    st.markdown("---")
    st.markdown("<div style='font-size:0.72rem;color:#6B7080;text-transform:uppercase;letter-spacing:0.08em;margin-bottom:8px;font-weight:600;'>🎓 AI Writing Coach</div>", unsafe_allow_html=True)
    if st.button("✨ Analyse My Writing", use_container_width=True, type="primary", key="coach_btn"):
        st.session_state.pop("coach_result", None)
        with st.spinner("Reading your prose..."):
            raw = _call_full(_coach_prompt(story, characters), max_tokens=320)
        st.session_state["coach_result"] = _parse_coach(raw)
        st.rerun()

    coach = st.session_state.get("coach_result")
    if coach:
        section_icons = {"STRENGTHS": ("💪", "#4D6BFE"), "WEAKNESSES": ("🔧", "#f87171"),
                         "SUGGESTIONS": ("💡", "#fbbf24"), "NEXT_SCENE": ("➡️", "#93c5fd")}
        for key, (icon, color) in section_icons.items():
            val = coach.get(key, "").strip()
            if val:
                st.markdown(f"""
                    <div class='issue-card' style='border-left-color:{color};margin-bottom:6px;'>
                        <div style='font-size:0.75rem;color:{color};font-weight:700;margin-bottom:4px;'>{icon} {key.replace("_"," ").title()}</div>
                        <div style='font-size:0.78rem;color:var(--text-primary);line-height:1.5;'>{_html.escape(val)}</div>
                    </div>
                """, unsafe_allow_html=True)


# ── World Notes sidebar ────────────────────────────────────────────────────────
def _sidebar_notes(story):
    username = st.session_state.username
    notes    = _cached_notes(username, story["id"])

    # Group notes by category
    grouped = {}
    for n in notes:
        grouped.setdefault(n["category"], []).append(n)

    if grouped:
        for cat, cat_notes in grouped.items():
            st.markdown(f"<div style='font-size:0.72rem;color:#4D6BFE;text-transform:uppercase;"
                        f"letter-spacing:0.1em;font-weight:700;margin:10px 0 4px;'>"
                        f"📂 {_html.escape(cat)}</div>", unsafe_allow_html=True)
            for n in cat_notes:
                st.markdown(f"""
                    <div class='char-card' style='margin-bottom:4px;'>
                        <div class='char-name'>{_html.escape(n['title'])}</div>
                        <div class='char-desc'>{_html.escape(n['content'][:140])}{'…' if len(n['content']) > 140 else ''}</div>
                    </div>
                """, unsafe_allow_html=True)
                nc1, nc2 = st.columns(2)
                with nc1:
                    if st.button("✏️ Edit", key=f"editnote_{n['id']}", use_container_width=True):
                        if st.session_state.get("editing_note") == n["id"]:
                            st.session_state.pop("editing_note", None)
                        else:
                            st.session_state["editing_note"] = n["id"]
                        st.rerun()
                with nc2:
                    if st.button("🗑️", key=f"delnote_{n['id']}", use_container_width=True):
                        delete_note(n["id"], username)
                        st.session_state.pop("editing_note", None)
                        _bust_cache()
                        st.rerun()
                if st.session_state.get("editing_note") == n["id"]:
                    en_cat  = st.selectbox("Category", NOTE_CATEGORIES,
                                index=NOTE_CATEGORIES.index(n["category"]) if n["category"] in NOTE_CATEGORIES else 0,
                                key=f"encat_{n['id']}")
                    en_ttl  = st.text_input("Title", value=n["title"],  key=f"entitle_{n['id']}")
                    en_body = st.text_area("Content", value=n["content"], height=80, key=f"enbody_{n['id']}")
                    ns1, ns2 = st.columns(2)
                    with ns1:
                        if st.button("💾 Save", key=f"savenote_{n['id']}", use_container_width=True):
                            update_note(n["id"], username, en_cat, en_ttl, en_body)
                            st.session_state.pop("editing_note", None)
                            _bust_cache()
                            st.rerun()
                    with ns2:
                        if st.button("✖ Cancel", key=f"cancelnote_{n['id']}", use_container_width=True):
                            st.session_state.pop("editing_note", None)
                            _bust_cache()
                            st.rerun()
    else:
        st.markdown("<div style='color:#6B7080;font-size:0.82rem;text-align:center;padding:16px;'>"
                    "No world notes yet.<br>Document your lore, magic, and geography here.</div>",
                    unsafe_allow_html=True)

    st.markdown("---")
    with st.expander("➕ Add Note"):
        n_cat  = st.selectbox("Category", NOTE_CATEGORIES, key="new_note_cat")
        n_ttl  = st.text_input("Title", key="new_note_title", placeholder="e.g. The Arcane Laws")
        n_body = st.text_area("Content", key="new_note_body", height=80,
                              placeholder="Describe your world detail...")
        if st.button("💾 Save Note", use_container_width=True, key="save_new_note"):
            if n_ttl.strip():
                add_note(username, story["id"], n_cat, n_ttl, n_body)
                for k in ["new_note_cat", "new_note_title", "new_note_body"]:
                    st.session_state.pop(k, None)
                _bust_cache()
                st.rerun()


# ── Snapshots sidebar ──────────────────────────────────────────────────────────
def _sidebar_snapshots(story):
    username  = st.session_state.username
    snapshots = _cached_snapshots(username, story["id"])

    if snapshots:
        st.markdown("<div style='color:#6B7080;font-size:0.78rem;margin-bottom:8px;'>"
                    "Restore any snapshot to roll back your story to that point.</div>",
                    unsafe_allow_html=True)
        for snap in snapshots:
            dt    = snap["created_at"][:16].replace("T", " ")
            st.markdown(f"""
                <div class='scene-card' style='margin-bottom:4px;'>
                    <div class='scene-title'>{_html.escape(snap['name'])}</div>
                    <div class='scene-meta'>📅 {_html.escape(dt)} · ✍️ {snap['word_count']:,} words</div>
                </div>
            """, unsafe_allow_html=True)
            sc1, sc2 = st.columns(2)
            with sc1:
                if st.button("Restore", key=f"restore_{snap['id']}", use_container_width=True):
                    msgs = restore_snapshot(snap["id"], username)
                    if msgs is not None:
                        story["messages"] = msgs
                        save_story(username, story)
                        st.session_state.show_cowrite_options = False
                        st.success(f"Restored: {snap['name']}")
                        _bust_cache()
                        st.rerun()
                    else:
                        # FIX #6: Report restore failure clearly
                        st.error("Restore failed — snapshot data may be corrupted or was deleted.")
            with sc2:
                if st.button("🗑️", key=f"delsnap_{snap['id']}", use_container_width=True):
                    delete_snapshot(snap["id"], username)
                    _bust_cache()
                    st.rerun()
    else:
        st.markdown("<div style='color:#6B7080;font-size:0.82rem;text-align:center;padding:16px;'>"
                    "No snapshots yet.<br>Save a snapshot before a risky edit.</div>",
                    unsafe_allow_html=True)

    st.markdown("---")
    if story.get("messages"):
        snap_name = st.text_input("Snapshot name", key="snap_name",
                                  placeholder=f"e.g. Before Chapter 3 rewrite")
        if st.button("📸 Save Snapshot", use_container_width=True, type="primary", key="save_snap"):
            name = snap_name.strip() or f"Snapshot {len(snapshots)+1}"
            save_snapshot(username, story["id"], name, story["messages"])
            st.session_state.pop("snap_name", None)
            st.success(f"Saved: {name}")
            _bust_cache()
            st.rerun()
    else:
        st.caption("Write some story first to create a snapshot.")


# ── Chapters sidebar ──────────────────────────────────────────────────────────
def _sidebar_chapters(story):
    username = st.session_state.username
    chapters = _cached_chapters(username, story["id"])
    scenes   = _cached_scenes(username, story["id"])

    if chapters:
        for ch in chapters:
            # Build list of scenes in this chapter
            ch_scenes = [sc for sc in scenes if sc.get("chapter_id") == ch["id"]]
            scene_list = ", ".join(sc["title"] for sc in ch_scenes) if ch_scenes else "No scenes assigned"
            st.markdown(f"""
                <div class='scene-card' style='border-left-color:rgba(77,107,254,0.50);'>
                    <div class='scene-title'>📖 Ch. {ch['order']} — {_html.escape(ch['title'])}</div>
                    <div class='scene-meta'>🎬 {_html.escape(scene_list)}</div>
                    {'<div class="scene-meta" style="font-style:italic;color:#6B7080;margin-top:4px;">' + _html.escape(ch["summary"][:80]) + ('…' if len(ch["summary"])>80 else '') + '</div>' if ch.get("summary") else ''}
                </div>
            """, unsafe_allow_html=True)

            ce1, ce2, ce3 = st.columns(3)
            with ce1:
                if st.button("✏️", key=f"edit_ch_{ch['id']}", use_container_width=True):
                    st.session_state["editing_chapter"] = (
                        None if st.session_state.get("editing_chapter") == ch["id"] else ch["id"])
                    st.rerun()
            with ce2:
                # AI-generate a chapter summary
                if ch_scenes and st.button("📝", key=f"sum_ch_{ch['id']}", use_container_width=True,
                                            help="AI-generate chapter summary"):
                    ai_msgs = [m for m in story.get("messages", []) if m["role"] == "assistant"]
                    excerpt = " ".join(m["content"] for m in ai_msgs)[:900]
                    prompt  = (f"Write a ONE sentence summary (max 25 words) of Chapter {ch['order']} "
                               f"titled '{ch['title']}' based on this prose:\n{excerpt}")
                    with st.spinner("Summarising…"):
                        summary = _call_full(prompt, max_tokens=50)
                    if summary:
                        save_chapter_summary(ch["id"], username, summary)
                        _bust_cache()
                        st.rerun()
            with ce3:
                if st.button("🗑️", key=f"del_ch_{ch['id']}", use_container_width=True):
                    delete_chapter(ch["id"], username)
                    _bust_cache()
                    st.rerun()

            # Edit form
            if st.session_state.get("editing_chapter") == ch["id"]:
                new_chtitle = st.text_input("Chapter title", value=ch["title"],
                                             key=f"chtitle_{ch['id']}")
                new_chsum   = st.text_area("Summary", value=ch.get("summary",""),
                                            height=60, key=f"chsum_{ch['id']}")
                # Scene assignment
                unassigned = [sc for sc in scenes
                              if sc.get("chapter_id") is None or sc.get("chapter_id") == ch["id"]]
                if unassigned:
                    st.markdown("<div style='font-size:0.68rem;color:#6B7080;margin:6px 0 2px;'>Assign scenes to this chapter:</div>",
                                unsafe_allow_html=True)
                    for sc in unassigned:
                        is_in = sc.get("chapter_id") == ch["id"]
                        if st.checkbox(sc["title"], value=is_in, key=f"assign_{sc['id']}_{ch['id']}"):
                            assign_scene_to_chapter(sc["id"], ch["id"], username)
                        elif is_in:
                            assign_scene_to_chapter(sc["id"], None, username)
                cs1, cs2 = st.columns(2)
                with cs1:
                    if st.button("💾 Save", key=f"save_ch_{ch['id']}", use_container_width=True):
                        update_chapter(ch["id"], username, new_chtitle, new_chsum)
                        st.session_state.pop("editing_chapter", None)
                        _bust_cache()
                        st.rerun()
                with cs2:
                    if st.button("✖ Cancel", key=f"cancel_ch_{ch['id']}", use_container_width=True):
                        st.session_state.pop("editing_chapter", None)
                        _bust_cache()
                        st.rerun()
    else:
        st.markdown("<div style='color:#6B7080;font-size:0.82rem;text-align:center;padding:16px;'>"
                    "No chapters yet.<br>Organise your story into chapters here.</div>",
                    unsafe_allow_html=True)

    st.markdown("---")
    new_ch_title = st.text_input("New chapter title", key="new_ch_title",
                                  placeholder="e.g. The Forest at Midnight")
    if st.button("➕ Add Chapter", use_container_width=True, type="primary", key="add_ch_btn"):
        title = new_ch_title.strip() or f"Chapter {len(chapters)+1}"
        add_chapter(username, story["id"], title)
        st.session_state.pop("new_ch_title", None)
        _bust_cache()
        st.rerun()


# ── Revision Mode ─────────────────────────────────────────────────────────────
def _revision_mode_panel(story, characters):
    """
    Paste any paragraph. Get 3 AI rewrites with different approaches.
    Shown below the chat history when the user activates it.
    """
    st.markdown("---")
    st.markdown(
        "<div style='font-size:0.72rem;font-weight:700;color:#6B7080;text-transform:uppercase;"
        "letter-spacing:0.14em;margin-bottom:10px;font-family:monospace;'>"
        "✂️ Revision Mode — Rewrite any paragraph</div>",
        unsafe_allow_html=True)

    rev_text = st.text_area(
        "Paste a paragraph to rewrite",
        height=110, key="revision_input",
        placeholder="Paste any paragraph from your story — the AI will generate 3 alternative versions…")

    REVISION_TONES = {
        "More literary / poetic":   "elevated, literary style with rich imagery and varied sentence rhythm",
        "Darker / more tense":      "dark, tense, foreboding tone with shorter punchy sentences",
        "Lighter / more hopeful":   "warmer, more hopeful and gentle tone with softer language",
        "Change to first person":   "first-person POV from the protagonist's perspective",
        "More concise (50% shorter)":"tighter, more concise prose — cut all filler, keep the essence",
        "Show don't tell":          "concrete sensory details instead of abstract statements",
    }
    selected_tones = st.multiselect(
        "Choose rewrite directions (pick 1–3)",
        list(REVISION_TONES.keys()),
        default=["More literary / poetic", "Darker / more tense", "More concise (50% shorter)"],
        key="revision_tones",
        max_selections=3,
    )

    if st.button("⚡ Generate Rewrites", use_container_width=True, type="primary", key="revision_btn"):
        if not rev_text.strip():
            st.warning("Paste a paragraph first.")
        elif not selected_tones:
            st.warning("Select at least one rewrite direction.")
        else:
            char_ctx = ""
            if characters:
                char_ctx = "Characters: " + "; ".join(
                    f"{c['name']} ({c['role']})" for c in characters) + "."
            rewrites = {}
            for tone_label in selected_tones[:3]:
                tone_dir = REVISION_TONES[tone_label]
                prompt = (
                    f"Rewrite the following paragraph in a {tone_dir}. "
                    f"Story genre: {story['genre']}, tone: {story['tone']}. "
                    f"{char_ctx} "
                    f"Return ONLY the rewritten paragraph. No preamble, no explanation.\n\n"
                    f"Original:\n{rev_text.strip()}"
                )
                with st.spinner(f"Generating: {tone_label}…"):
                    result = _call_full(prompt, max_tokens=220)
                rewrites[tone_label] = result

            st.session_state["revision_results"] = rewrites
            st.rerun()

    # Show results
    results = st.session_state.get("revision_results")
    if results:
        st.markdown("<div style='margin-top:12px;'></div>", unsafe_allow_html=True)
        for label, text in results.items():
            if not text:
                continue
            st.markdown(f"""
                <div style='background:var(--bg-card);border:1px solid rgba(77,107,254,0.18);
                border-left:3px solid #4D6BFE;border-radius:10px;padding:16px 18px;margin:8px 0;'>
                    <div style='font-size:0.65rem;font-weight:700;color:#4D6BFE;letter-spacing:0.12em;
                    text-transform:uppercase;font-family:monospace;margin-bottom:8px;'>✏️ {_html.escape(label)}</div>
                    <div style='color:var(--text-primary);font-size:1rem;line-height:1.75;
                    font-family:'Inter',sans-serif;'>{_html.escape(text)}</div>
                </div>
            """, unsafe_allow_html=True)
            if st.button(f"➕ Use this version", key=f"use_rev_{hash(label)}", use_container_width=True):
                story["messages"].append({"role": "assistant", "content": f"[Revision — {label}]\n\n{text}"})
                save_story(st.session_state.username, story)
                st.session_state.pop("revision_results", None)
                st.session_state.pop("revision_input",  None)
                _bust_cache()
                st.rerun()
        if st.button("🗑️ Clear Rewrites", use_container_width=True, key="clear_revisions"):
            st.session_state.pop("revision_results", None)
            _bust_cache()
            st.rerun()


# ══════════════════════════════════════════════════════════════════════════════
# INTELLIGENCE LAYER — 10 advanced analysis features
# ══════════════════════════════════════════════════════════════════════════════

# ── Show Don't Tell — static detector ────────────────────────────────────────
_TELL_PATTERNS = [
    # Emotion tells
    (r'\b(was|felt|seemed?|appeared?|looked?)\s+(very\s+)?(angry|sad|happy|scared|afraid|nervous|excited|confused|surprised|shocked|embarrassed|guilty|proud|jealous|lonely|tired|bored|anxious|frustrated)\b',
     "Emotion tell — show the physical reaction instead"),
    (r'\bhe was (angry|sad|happy|scared|afraid)\b',
     "Weak emotion tell — try showing body language"),
    (r'\bshe was (angry|sad|happy|scared|afraid)\b',
     "Weak emotion tell — try showing body language"),
    # Adverb overuse
    (r'\b\w+ly\s+(said|asked|replied|answered|shouted|whispered|cried|laughed)\b',
     "Adverb + dialogue verb — let the dialogue speak"),
    # Abstract states
    (r'\b(it was|there was|there were)\s+(a\s+)?(beautiful|amazing|terrible|horrible|wonderful|great|awful)\b',
     "Abstract descriptor — use a concrete image instead"),
    # Weak openers
    (r'^(It was|There was|There were|He was|She was|They were)\b',
     "Weak sentence opener — start with action or image"),
    # Passive voice indicators
    (r'\b(was|were|been|is|are)\s+\w+ed\b',
     "Possible passive voice — consider active construction"),
    # Filler phrases
    (r'\b(suddenly|all of a sudden|out of nowhere|without warning)\b',
     "Filler word — build tension through pacing instead"),
    (r'\b(very|really|quite|rather|somewhat|fairly|pretty)\s+\w+\b',
     "Intensifier weakens prose — try a stronger word"),
]

def _show_dont_tell(prose):
    """Returns list of {sentence, issue, suggestion} dicts."""
    found = []
    for sent in _sentences(prose):
        for pattern, tip in _TELL_PATTERNS:
            if re.search(pattern, sent, re.IGNORECASE):
                found.append({"sentence": sent[:120], "issue": tip})
                break  # one flag per sentence
    return found[:10]


# ── Tension Meter ─────────────────────────────────────────────────────────────
_TENSION_HIGH = {
    "death","kill","killed","blood","scream","screamed","shook","trembled",
    "danger","trap","trapped","escape","chase","chased","fight","fighting",
    "heart pounding","sweat","gasped","panic","fear","terror","threat",
    "weapon","blade","fire","burning","crash","explosion","attack","flee",
    "desperate","helpless","wound","wounded","dying","monster","shadow",
    "dark","darkness","alone","lost","betrayal","betrayed","lied","lie",
}
_TENSION_LOW = {
    "smiled","laughed","peace","calm","gentle","soft","quiet","warm",
    "safe","home","sunlight","morning","rested","breathed","sighed","relief",
    "beautiful","wonder","joy","happy","content","comfortable","familiar",
}

def _tension_meter(story):
    """Score each AI message for tension 0–100."""
    ai_msgs = [m for m in story.get("messages", [])
               if m["role"] == "assistant" and not m["content"].startswith("◆")]
    results = []
    for i, msg in enumerate(ai_msgs):
        words = re.sub(r'[^\w\s]', '', msg["content"].lower()).split()
        hi = sum(1 for w in words if w in _TENSION_HIGH)
        lo = sum(1 for w in words if w in _TENSION_LOW)
        total = max(1, hi + lo)
        score = int((hi / total) * 100)
        results.append({"label": f"Block {i+1}", "score": score,
                         "words": len(words)})
    return results


# ── Plot Hole Detector — AI powered ───────────────────────────────────────────
def _plot_hole_prompt(story, characters, scenes):
    prose   = _all_prose(story)[:2000]
    char_names = [c["name"] for c in characters] if characters else []
    scene_titles = [sc["title"] for sc in scenes] if scenes else []
    arc     = story.get("plot_arc", {})
    arc_done = [label for key, label in PLOT_STAGES if arc.get(key)]

    return (
        "You are a professional story editor. Analyse the following story excerpt for "
        "narrative problems. Find EXACTLY the following types of issues:\n\n"
        "UNRESOLVED_THREADS: (things introduced but never followed up)\n"
        "CONTRADICTIONS: (facts that conflict with each other)\n"
        "MISSING_MOTIVATION: (characters acting without clear reason)\n"
        "TIMELINE_GAPS: (events out of logical order or unexplained jumps)\n\n"
        f"Characters: {', '.join(char_names) if char_names else 'none defined'}\n"
        f"Scenes: {', '.join(scene_titles) if scene_titles else 'none defined'}\n"
        f"Arc completed: {', '.join(arc_done) if arc_done else 'not started'}\n\n"
        f"Story excerpt:\n{prose}\n\n"
        "Be specific — reference actual text. If no issues found in a category, write 'None found.' "
        "Keep total response under 220 words."
    )

def _parse_plot_holes(raw):
    categories = ["UNRESOLVED_THREADS", "CONTRADICTIONS",
                  "MISSING_MOTIVATION", "TIMELINE_GAPS"]
    result = {c: "" for c in categories}
    current = None
    for line in raw.splitlines():
        s = line.strip()
        for cat in categories:
            if s.upper().startswith(cat + ":") or s.upper() == cat:
                current = cat
                tail = s[len(cat):].lstrip(": ").strip()
                if tail:
                    result[cat] += tail + "\n"
                break
        else:
            if current and s:
                result[current] += s + "\n"
    return {k: v.strip() for k, v in result.items()}


# ── Character Voice Checker — AI powered ──────────────────────────────────────
def _voice_check_prompt(story, characters):
    if not characters:
        return None
    prose = _all_prose(story)[:1800]
    char_block = "\n".join(
        f"- {c['name']} ({c['role']}): speaking style = '{c.get('speaking_style') or 'not defined'}'"
        for c in characters
    )
    return (
        "You are a dialogue coach and story editor. Analyse whether each character "
        "speaks in a distinct, consistent voice in the following story.\n\n"
        f"Characters and their intended speaking styles:\n{char_block}\n\n"
        f"Story prose:\n{prose}\n\n"
        "For each character, write one line:\n"
        "CHARACTER_NAME: [score 1-10] — [one sentence on what works or what to fix]\n\n"
        "Then write:\n"
        "OVERALL: [one sentence summary of voice consistency across the cast]\n\n"
        "Be direct and specific. Under 150 words total."
    )


# ── Story Timeline ────────────────────────────────────────────────────────────
def _render_timeline(chapters, scenes):
    """Render a visual timeline using scenes grouped into chapters."""
    if not chapters and not scenes:
        st.markdown(
            "<div style='color:#6B7080;font-size:0.82rem;text-align:center;padding:20px;'>"
            "Add chapters and scenes to see your story timeline.</div>",
            unsafe_allow_html=True)
        return

    # Group scenes by chapter
    ungrouped = [sc for sc in scenes if not sc.get("chapter_id")]
    ch_map    = {ch["id"]: ch for ch in chapters}

    html_parts = ["<div style='padding:8px 0;'>"]

    for ch in chapters:
        ch_scenes = [sc for sc in scenes if sc.get("chapter_id") == ch["id"]]
        html_parts.append(f"""
            <div style='margin-bottom:20px;'>
                <div style='display:flex;align-items:center;gap:10px;margin-bottom:8px;'>
                    <div style='width:32px;height:32px;border-radius:50%;
                        background:linear-gradient(135deg,#4D6BFE,#B8943F);
                        display:flex;align-items:center;justify-content:center;
                        font-size:0.75rem;font-weight:800;color:#060e06;flex-shrink:0;'>
                        {ch['order']}
                    </div>
                    <div>
                        <div style='font-weight:700;color:var(--text-primary);font-size:0.92rem;'>
                            {_html.escape(ch['title'])}</div>
                        {'<div style="font-size:0.72rem;color:#6B7080;font-style:italic;">' + _html.escape(ch['summary'][:60]) + '</div>' if ch.get('summary') else ''}
                    </div>
                </div>
        """)
        if ch_scenes:
            html_parts.append("<div style='margin-left:42px;border-left:2px solid rgba(77,107,254,0.20);padding-left:14px;'>")
            for sc in ch_scenes:
                html_parts.append(f"""
                    <div style='display:flex;gap:8px;align-items:flex-start;margin-bottom:8px;position:relative;'>
                        <div style='width:8px;height:8px;border-radius:50%;background:#4D6BFE;
                            margin-top:5px;flex-shrink:0;box-shadow:0 0 6px rgba(77,107,254,0.50);
                            position:absolute;left:-18px;'></div>
                        <div style='background:var(--bg-card);border:1px solid rgba(77,107,254,0.12);
                            border-radius:8px;padding:8px 12px;flex:1;'>
                            <div style='font-size:0.82rem;font-weight:600;color:var(--text-primary);'>
                                🎬 {_html.escape(sc['title'])}</div>
                            {'<div style="font-size:0.7rem;color:#6B7080;margin-top:2px;">📍 ' + _html.escape(sc['location']) + '</div>' if sc.get('location') else ''}
                            {'<div style="font-size:0.7rem;color:#6B7080;">👥 ' + _html.escape(', '.join(sc['characters'][:3])) + '</div>' if sc.get('characters') else ''}
                        </div>
                    </div>
                """)
            html_parts.append("</div>")
        else:
            html_parts.append(
                "<div style='margin-left:42px;font-size:0.72rem;color:#6B7080;"
                "font-style:italic;padding:4px 0 8px;'>No scenes assigned yet</div>")
        html_parts.append("</div>")

    if ungrouped:
        html_parts.append("""
            <div style='margin-top:12px;border-top:1px solid rgba(77,107,254,0.08);padding-top:12px;'>
                <div style='font-size:0.7rem;color:#6B7080;font-weight:700;text-transform:uppercase;
                    letter-spacing:0.1em;margin-bottom:8px;'>Unassigned Scenes</div>
        """)
        for sc in ungrouped:
            html_parts.append(f"""
                <div style='background:var(--bg-surface);border:1px solid rgba(77,107,254,0.08);
                    border-radius:8px;padding:8px 12px;margin-bottom:6px;'>
                    <div style='font-size:0.8rem;color:#A8BCFF;'>🎬 {_html.escape(sc['title'])}</div>
                    {'<div style="font-size:0.68rem;color:#6B7080;">📍 ' + _html.escape(sc['location']) + '</div>' if sc.get('location') else ''}
                </div>
            """)
        html_parts.append("</div>")

    html_parts.append("</div>")
    st.markdown("".join(html_parts), unsafe_allow_html=True)


# ── Reading Mode (Premium — 5 themes, bookmarks, page nav) ───────────────────
def show_reading_mode():
    """Premium book reading experience — delegates to reading_mode.py."""
    story = _get_story()
    if not story:
        st.error("No story loaded.")
        return
    if _HAS_READING_MODE:
        show_premium_reading_mode(story)
    else:
        _reading_mode_fallback(story)

def _reading_mode_fallback(story):
    """Basic fallback reading mode (used if reading_mode.py is absent)."""
    ai_msgs = [m for m in story.get("messages", [])
               if m["role"] == "assistant"
               and not m["content"].startswith("◆")]

    # Inject reading mode CSS
    st.markdown("""
        <style>
        .reading-body {
            max-width: 680px; margin: 0 auto;
            font-family: 'Inter', sans-serif;
            font-size: 1.18rem; line-height: 1.95;
            color:var(--text-primary); padding: 60px 24px 120px;
        }
        .reading-title {
            font-size: clamp(2rem, 5vw, 3.2rem);
            font-weight: 700; color: #4D6BFE;
            text-align: center; margin-bottom: 6px;
            letter-spacing: -0.01em;
        }
        .reading-meta {
            text-align: center; color: #6B7080;
            font-size: 0.82rem; letter-spacing: 0.12em;
            text-transform: uppercase; margin-bottom: 60px;
            font-family: 'JetBrains Mono', monospace;
        }
        .reading-para {
            margin-bottom: 1.6em;
            text-indent: 2em;
        }
        .reading-para:first-of-type { text-indent: 0; }
        .reading-divider {
            text-align: center; color: rgba(77,107,254,0.30);
            margin: 48px 0; font-size: 1.2rem;
            letter-spacing: 0.4em;
        }
        </style>
    """, unsafe_allow_html=True)

    # Back button
    col1, col2, col3 = st.columns([1, 4, 1])
    with col1:
        if st.button("← Back to Writing", key="reading_back"):
            st.session_state["reading_mode"] = False
            st.rerun()
    with col3:
        wc = _word_count(story)
        st.markdown(
            f"<div style='text-align:right;font-size:0.72rem;color:#6B7080;"
            f"font-family:monospace;padding-top:8px;'>{wc:,} words · {_reading_time(story)}</div>",
            unsafe_allow_html=True)

    if not ai_msgs:
        st.markdown(
            "<div style='text-align:center;color:#6B7080;padding:80px;font-style:italic;'>"
            "Nothing written yet. Go back and start your story.</div>",
            unsafe_allow_html=True)
        return

    # Assemble prose
    full_prose = "\n\n".join(m["content"] for m in ai_msgs)
    paragraphs = [p.strip() for p in full_prose.split("\n\n") if p.strip()]

    html = [
        f"<div class='reading-body'>",
        f"<div class='reading-title'>{_html.escape(story['title'])}</div>",
        f"<div class='reading-meta'>{_html.escape(story['genre'])} · {_html.escape(story['tone'])} · {_word_count(story):,} words</div>",
    ]
    for i, para in enumerate(paragraphs):
        if para.startswith("[Revision"):
            html.append(f"<div class='reading-divider'>✦ ✦ ✦</div>")
            continue
        html.append(f"<p class='reading-para'>{_html.escape(para)}</p>")
    html.append("</div>")

    st.markdown("".join(html), unsafe_allow_html=True)


# ── PDF Export ────────────────────────────────────────────────────────────────
def _export_pdf(story, characters, scenes):
    """
    Generate a formatted PDF as bytes using HTML→reportlab.
    Falls back to a clean plain-text PDF if reportlab not available.
    """
    prose = " ".join(m["content"] for m in story.get("messages", [])
                     if m["role"] == "assistant"
                     and not m["content"].startswith("◆"))

    # Build an HTML string then convert via weasyprint if available,
    # otherwise produce a minimal reportlab PDF
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import (SimpleDocTemplate, Paragraph,
                                         Spacer, HRFlowable)
        from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY

        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4,
                                 leftMargin=3*cm, rightMargin=3*cm,
                                 topMargin=3*cm,  bottomMargin=3*cm)

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle("NF_Title",
            fontSize=28, leading=34, alignment=TA_CENTER,
            textColor=colors.HexColor("#1a3a1a"),
            fontName="Times-Bold", spaceAfter=6)
        meta_style = ParagraphStyle("NF_Meta",
            fontSize=10, alignment=TA_CENTER,
            textColor=colors.HexColor("#6B7080"),
            fontName="Times-Italic", spaceAfter=24)
        body_style = ParagraphStyle("NF_Body",
            fontSize=12, leading=20, alignment=TA_JUSTIFY,
            textColor=colors.HexColor("#1a1a1a"),
            fontName="Times-Roman", firstLineIndent=24,
            spaceAfter=10)
        char_head = ParagraphStyle("NF_CharHead",
            fontSize=11, leading=14,
            textColor=colors.HexColor("#1a3a1a"),
            fontName="Times-Bold", spaceBefore=8, spaceAfter=2)
        char_body = ParagraphStyle("NF_CharBody",
            fontSize=10, leading=14,
            textColor=colors.HexColor("#444444"),
            fontName="Times-Italic", spaceAfter=6)

        elems = []
        elems.append(Paragraph(story["title"], title_style))
        elems.append(Paragraph(
            f"{story['genre']} · {story['tone']} · {_word_count(story):,} words",
            meta_style))
        elems.append(HRFlowable(width="100%", thickness=0.5,
                                  color=colors.HexColor("#4D6BFE")))
        elems.append(Spacer(1, 0.4*cm))

        if characters:
            elems.append(Paragraph("Characters", styles["Heading2"]))
            for c in characters:
                elems.append(Paragraph(f"{c['name']} — {c['role']}", char_head))
                if c.get("description"):
                    elems.append(Paragraph(c["description"], char_body))
            elems.append(HRFlowable(width="100%", thickness=0.3,
                                      color=colors.HexColor("#cccccc")))
            elems.append(Spacer(1, 0.4*cm))

        elems.append(Paragraph("Story", styles["Heading2"]))
        elems.append(Spacer(1, 0.2*cm))
        for para in prose.split("\n\n"):
            para = para.strip()
            if para:
                safe = para.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
                elems.append(Paragraph(safe, body_style))

        doc.build(elems)
        return buf.getvalue()

    except ImportError:
        # FIX #8: Do NOT return a fake .txt disguised as .pdf
        # Raise so the download button shows an error instead of a corrupt file
        raise ImportError(
            "reportlab is required for PDF export. "
            "Run: pip install reportlab==4.2.2"
        )


# ── Intelligence Panel (new sidebar tab) ─────────────────────────────────────
# ══════════════════════════════════════════════════════════════════════════════
#  NEW FEATURES v4
# ══════════════════════════════════════════════════════════════════════════════

# ── Feature 1: Clickable Title Rename ─────────────────────────────────────────
def _apply_title(story, title_line):
    """Strip numbering and apply a title from the generator to the story."""
    import re as _re
    clean = _re.sub(r"^\d+[\.\)]\s*", "", title_line).strip().strip('"').strip("'")
    if clean:
        story["title"] = clean
        from database import update_story_title
        update_story_title(st.session_state.username, story["id"], clean)


# ── Feature 2: Reading Level Analyzer ─────────────────────────────────────────
def _count_syllables(word):
    word = word.lower().strip(".,!?;:")
    if len(word) <= 3:
        return 1
    count = 0
    vowels = "aeiouy"
    prev_vowel = False
    for ch in word:
        is_v = ch in vowels
        if is_v and not prev_vowel:
            count += 1
        prev_vowel = is_v
    if word.endswith("e"):
        count -= 1
    return max(1, count)

def _reading_level(text):
    """Returns Flesch-Kincaid grade, reading ease, and label."""
    words = text.split()
    if len(words) < 30:
        return None
    sentences = max(1, len([s for s in re.split(r'[.!?]+', text) if s.strip()]))
    syllables  = sum(_count_syllables(w) for w in words)
    asl = len(words) / sentences          # avg sentence length
    asw = syllables / len(words)          # avg syllables per word
    ease  = 206.835 - 1.015 * asl - 84.6 * asw
    grade = 0.39 * asl + 11.8 * asw - 15.59
    ease  = max(0, min(100, ease))
    grade = max(1, grade)
    if ease >= 80:   label = "Very Easy (Children)"
    elif ease >= 70: label = "Easy (Young Adult)"
    elif ease >= 60: label = "Standard (General)"
    elif ease >= 50: label = "Fairly Difficult (Academic)"
    elif ease >= 30: label = "Difficult (Literary)"
    else:            label = "Very Difficult (Scholarly)"
    return {"ease": round(ease, 1), "grade": round(grade, 1), "label": label,
            "words": len(words), "sentences": sentences, "syllables": syllables}


# ── Feature 3: Cliché Detector ────────────────────────────────────────────────
def _detect_cliches(text):
    t = text.lower()
    found = []
    for c in _CLICHES:
        if c in t:
            found.append(c)
    return found


# ── Feature 4: Sentiment Arc Chart Data ───────────────────────────────────────
def _sentiment_arc(story):
    """Split story prose into chunks, score each, return list of (label, score)."""
    all_msgs = [m["content"] for m in story.get("messages", [])
                if not m["content"].startswith("◆")]
    if not all_msgs:
        return []
    points = []
    for i, msg in enumerate(all_msgs):
        words = msg.lower().split()
        pos = sum(1 for w in words if w in _POS_WORDS)
        neg = sum(1 for w in words if w in _NEG_WORDS)
        total = max(1, len(words))
        score = round(((pos - neg) / total) * 100, 1)
        points.append({"Message": i + 1, "Sentiment": score})
    return points


# ── Feature 5: Character Interview ────────────────────────────────────────────
def _sidebar_interview(story, characters):
    st.markdown(
        "<div style='font-size:0.7rem;color:var(--text-muted);margin-bottom:10px;'>"
        "Ask your character anything. They'll answer <em>in their own voice</em>.</div>",
        unsafe_allow_html=True)
    if not characters:
        st.markdown(
            "<div style='color:var(--text-muted);font-size:0.82rem;text-align:center;padding:20px;'>"
            "Add characters in the Cast tab first.</div>", unsafe_allow_html=True)
        return

    char_names = [c["name"] for c in characters]
    sel_name   = st.selectbox("Speak with", char_names, key="interview_char",
                               label_visibility="collapsed")
    char       = next((c for c in characters if c["name"] == sel_name), None)

    # Show character card
    if char:
        st.markdown(
            f"<div class='char-card' style='margin-bottom:10px;'>"
            f"<div style='font-weight:600;font-size:0.85rem;color:var(--text-primary);'>"
            f"{_html.escape(char['name'])} "
            f"<span style='font-size:0.72rem;color:var(--primary);font-weight:400;'>"
            f"({_html.escape(char.get('role',''))})</span></div>"
            f"<div style='font-size:0.75rem;color:var(--text-muted);margin-top:3px;'>"
            f"{_html.escape(char.get('description','')[:120])}</div>"
            f"</div>",
            unsafe_allow_html=True)

    question = st.text_input(
        "Your question", placeholder="What drives you? What are you afraid of?",
        key="interview_q", max_chars=200, label_visibility="collapsed")

    if st.button("Ask", use_container_width=True, key="interview_ask", type="primary"):
        if question.strip() and char:
            prose_snip = _all_prose(story)[:500]
            prompt = (
                f"You are {char['name']}, a {char.get('role','character')} in a "
                f"{story['genre']} story.\n"
                f"Your description: {char.get('description','')}\n"
                f"Your speaking style: {char.get('speaking_style','natural and distinct')}\n"
                f"Story context: {prose_snip or 'The story is just beginning.'}\n\n"
                f"RULES:\n"
                f"1. Answer ONLY as {char['name']} — first person, in character.\n"
                f"2. Reflect your personality, fears, and speaking style.\n"
                f"3. Do NOT break character or add any meta-commentary.\n"
                f"4. Keep the answer to 3-5 sentences.\n\n"
                f"Question: {question}"
            )
            with st.spinner(f"{char['name']} is thinking…"):
                answer = _call_full(prompt, max_tokens=200)
            if answer == llm.AI_UNAVAILABLE:
                st.error("⚠️ AI unavailable — is Ollama running?")
                st.stop()
            # Store Q&A history
            history = st.session_state.get(f"interview_history_{story['id']}", [])
            history.append({"char": char["name"], "q": question, "a": answer})
            st.session_state[f"interview_history_{story['id']}"] = history[-10:]
            st.rerun()  # interview_q cleared via callback

    # Display interview history
    history = st.session_state.get(f"interview_history_{story['id']}", [])
    if history:
        st.markdown("<div style='margin-top:12px;'></div>", unsafe_allow_html=True)
        for item in reversed(history):
            st.markdown(
                f"<div class='bubble-user' style='margin-bottom:4px;max-width:100%;margin-left:0;'>"
                f"<div class='lbl'>You → {_html.escape(item['char'])}</div>"
                f"<div class='txt'>{_html.escape(item['q'])}</div></div>",
                unsafe_allow_html=True)
            st.markdown(
                f"<div class='bubble-ai' style='margin-bottom:8px;'>"
                f"<div class='lbl'>{_html.escape(item['char'])}</div>"
                f"<div class='txt'>{_html.escape(item['a'])}</div></div>",
                unsafe_allow_html=True)
        if st.button("🗑 Clear Interview", key="clear_interview", use_container_width=True):
            st.session_state.pop(f"interview_history_{story['id']}", None)
            st.rerun()


# ── Feature 6: Story Bible Export ─────────────────────────────────────────────
def _export_story_bible(story, characters, scenes, username):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document()

    def _heading(doc, text, level=1, color=(45, 85, 232)):
        p = doc.add_heading(text, level=level)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
        return p

    def _para(doc, text, italic=False, bold=False):
        p = doc.add_paragraph(text)
        for run in p.runs:
            run.italic = italic
            run.bold   = bold
        return p

    # Title page
    t = doc.add_heading(story["title"], 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(f"Genre: {story['genre']}   ·   Tone: {story['tone']}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Author: {username}")
    doc.add_page_break()

    # Story summary
    _heading(doc, "Story Summary")
    prose = _all_prose(story)[:1200]
    if prose:
        _para(doc, prose + ("…" if len(prose) == 1200 else ""), italic=True)
    else:
        _para(doc, "No story content yet.", italic=True)
    doc.add_paragraph()

    # Stats
    _heading(doc, "Quick Stats", 2)
    wc = _word_count(story)
    doc.add_paragraph(f"• Word count: {wc:,}")
    doc.add_paragraph(f"• Messages: {len(story.get('messages', []))}")
    doc.add_paragraph(f"• Estimated read time: ~{max(1, wc//200)} minutes")
    rl = _reading_level(prose) if prose else None
    if rl:
        doc.add_paragraph(f"• Reading level: {rl['label']} (Grade {rl['grade']})")
    doc.add_page_break()

    # Characters
    _heading(doc, "Characters")
    if characters:
        for c in characters:
            _heading(doc, c["name"], 2)
            doc.add_paragraph(f"Role: {c.get('role', '—')}")
            doc.add_paragraph(f"Description: {c.get('description', '—')}")
            if c.get("speaking_style"):
                doc.add_paragraph(f"Speaking style: {c['speaking_style']}")
            doc.add_paragraph()
    else:
        _para(doc, "No characters added yet.", italic=True)
    doc.add_page_break()

    # Scenes
    _heading(doc, "Scenes")
    if scenes:
        for i, sc in enumerate(scenes, 1):
            _heading(doc, f"Scene {i}: {sc.get('title','Untitled')}", 2)
            doc.add_paragraph(f"Location: {sc.get('location', '—')}")
            doc.add_paragraph(f"Purpose: {sc.get('purpose', '—')}")
            if sc.get("description"):
                _para(doc, sc["description"], italic=True)
            doc.add_paragraph()
    else:
        _para(doc, "No scenes added yet.", italic=True)
    doc.add_page_break()

    # Plot arc
    _heading(doc, "Plot Arc")
    arc = story.get("plot_arc", {})
    for key, label in PLOT_STAGES:
        _heading(doc, label, 2)
        _para(doc, arc.get(key, "Not written yet."), italic=not arc.get(key))
        doc.add_paragraph()
    doc.add_page_break()

    # Full story text
    _heading(doc, "Full Story Text")
    msgs = story.get("messages", [])
    if msgs:
        for m in msgs:
            if m["role"] == "assistant" and not m["content"].startswith("◆"):
                doc.add_paragraph(m["content"])
                doc.add_paragraph()
    else:
        _para(doc, "No story text yet.", italic=True)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ══════════════════════════════════════════════════════════════════════════════
#  FINAL FEATURES v5
# ══════════════════════════════════════════════════════════════════════════════

# ── Genre Craft Tips ──────────────────────────────────────────────────────────
_GENRE_TIPS = {
    "Fantasy": [
        "🌍 Build the rules of your world before you break them — readers trust consistency.",
        "⚔️ Ground magic in cost or consequence — power without price feels hollow.",
        "🗺️ Name locations and objects with phonetic logic — invented words should feel speakable.",
    ],
    "Sci-Fi": [
        "🔬 One speculative idea done deeply beats five ideas done shallowly.",
        "🤖 Technology should illuminate something about human nature, not replace it.",
        "📡 Anchor the unfamiliar in familiar sensory details — smell, texture, sound.",
    ],
    "Mystery": [
        "🔍 Plant every clue fairly — the reader should be able to solve it if sharp enough.",
        "🪤 Red herrings work only if they have their own story logic — not just misdirection.",
        "⏱️ Pace revelation carefully — one answer should always raise two new questions.",
    ],
    "Romance": [
        "💔 Tension lives in the gap between want and action — delay the obvious.",
        "👁️ Show attraction through specific, physical, sensory detail — not adjectives.",
        "🌧️ The best obstacle is internal, not external — fear, not circumstance.",
    ],
    "Horror": [
        "🕯️ What the reader imagines is always scarier than what you show.",
        "😰 Dread is built in the ordinary — make the mundane feel wrong.",
        "🔇 Silence and stillness are your most powerful tools. Use them.",
    ],
    "Thriller": [
        "⏰ Every scene needs a ticking clock — time pressure creates forward momentum.",
        "🎭 Your antagonist needs a logic the reader can follow, even if they can't agree.",
        "📌 End chapters on a question, not an answer.",
    ],
    "Adventure": [
        "🧭 The external journey should mirror an internal one — geography as metaphor.",
        "⚡ Vary pace deliberately — sprint scenes need walking scenes to contrast.",
        "🤝 Ensemble casts need distinct voices — each character should sound irreplaceable.",
    ],
    "Historical": [
        "📜 Period details work best when they surprise — avoid the obvious.",
        "🗣️ Dialogue should feel period-appropriate but remain readable — no phonetic spelling.",
        "🏛️ Let history be the backdrop, not the story — the story is always personal.",
    ],
}

def _show_genre_tips(genre):
    tips = _GENRE_TIPS.get(genre, [])
    if not tips:
        return
    st.markdown(
        f"<div style='margin-top:10px;'>"
        f"<div style='font-size:0.65rem;font-family:JetBrains Mono,monospace;color:var(--text-muted);"
        f"text-transform:uppercase;letter-spacing:0.1em;margin-bottom:6px;'>✦ {genre} Craft Tips</div>",
        unsafe_allow_html=True)
    for tip in tips:
        st.markdown(
            f"<div style='font-size:0.75rem;color:var(--text-primary);line-height:1.5;"
            f"padding:5px 8px;border-left:2px solid var(--primary-border);margin-bottom:4px;'>"
            f"{tip}</div>", unsafe_allow_html=True)
    st.markdown("</div>", unsafe_allow_html=True)


# ── Pomodoro Focus Timer ──────────────────────────────────────────────────────
def _pomodoro_widget(story_id):
    key_end   = f"_pomo_end_{story_id}"
    key_label = f"_pomo_label_{story_id}"
    key_dur   = f"_pomo_dur_{story_id}"

    now = time.time()
    end = st.session_state.get(key_end, 0)
    running = end > now

    st.markdown(
        "<div style='font-size:0.65rem;font-family:JetBrains Mono,monospace;"
        "color:var(--text-muted);text-transform:uppercase;letter-spacing:0.1em;"
        "margin-bottom:8px;'>⏱ Focus Timer</div>",
        unsafe_allow_html=True)

    if running:
        remaining = int(end - now)
        mins, secs = divmod(remaining, 60)
        label = st.session_state.get(key_label, "Focus")
        pct   = remaining / st.session_state.get(key_dur, 1500)
        bar_w = int((1 - pct) * 100)
        st.markdown(
            f"<div style='background:var(--bg-card);border:1px solid var(--primary-border);"
            f"border-radius:var(--radius-md);padding:10px 12px;text-align:center;'>"
            f"<div style='font-size:1.4rem;font-weight:700;color:var(--primary);"
            f"font-family:JetBrains Mono,monospace;letter-spacing:0.05em;'>"
            f"{mins:02d}:{secs:02d}</div>"
            f"<div style='font-size:0.65rem;color:var(--text-muted);margin-top:2px;'>{label}</div>"
            f"<div style='background:var(--primary-dim);border-radius:99px;height:3px;margin-top:8px;'>"
            f"<div style='background:var(--primary);width:{bar_w}%;height:3px;border-radius:99px;"
            f"transition:width 1s linear;'></div></div>"
            f"</div>", unsafe_allow_html=True)
        if st.button("⏹ Stop", key=f"pomo_stop_{story_id}", use_container_width=True):
            st.session_state.pop(key_end, None)
            st.rerun()
        # Nudge rerun every 10s via meta-refresh (safer than location.reload)
        import streamlit.components.v1 as _sc
        _sc.html(
            "<script>setTimeout(function(){"
            "var e=new CustomEvent('streamlit:rerequestrender',{});"
            "window.parent.dispatchEvent(e);"
            "},10000);</script>",
            height=0)
    else:
        if end != 0 and end <= now:
            st.success("🎉 Session complete! Take a short break.")

        dur_options = {"25 min (Pomodoro)": 1500, "15 min": 900,
                       "45 min": 2700, "5 min (Break)": 300}
        chosen = st.selectbox("Duration", list(dur_options.keys()),
                              key=f"pomo_dur_sel_{story_id}",
                              label_visibility="collapsed")
        label_input = st.text_input("Label", value="Focus session",
                                    key=f"pomo_label_{story_id}",
                                    label_visibility="collapsed",
                                    placeholder="What are you working on?")
        if st.button("▶ Start Timer", key=f"pomo_start_{story_id}",
                     use_container_width=True, type="primary"):
            dur = dur_options[chosen]
            st.session_state[key_end]   = now + dur
            st.session_state[key_dur]   = dur
            st.session_state[key_label] = label_input or "Focus session"
            st.rerun()


# ── Search In Story ────────────────────────────────────────────────────────────
def _sidebar_search_story(story):
    st.markdown(
        "<div style='font-size:0.7rem;color:var(--text-muted);margin-bottom:8px;'>"
        "Search through every message in this story.</div>",
        unsafe_allow_html=True)
    query = st.text_input("Search story", placeholder="🔍 Search…",
                          key="ws_search_q", label_visibility="collapsed",
                          max_chars=100)
    if not query.strip():
        total = len(story.get("messages", []))
        wc    = _word_count(story)
        st.markdown(
            f"<div style='font-size:0.75rem;color:var(--text-muted);padding:8px 0;'>"
            f"{total} messages · {wc:,} words total</div>",
            unsafe_allow_html=True)
        return

    q   = query.lower()
    hits = [(i, m) for i, m in enumerate(story.get("messages", []))
            if q in m["content"].lower()]

    if not hits:
        st.markdown(
            f"<div style='color:var(--text-muted);font-size:0.82rem;padding:12px 0;'>"
            f"No results for \"{_html.escape(query)}\"</div>",
            unsafe_allow_html=True)
        return

    st.markdown(
        f"<div style='font-size:0.72rem;color:var(--primary);margin-bottom:8px;font-weight:600;'>"
        f"{len(hits)} result(s) for \"{_html.escape(query)}\"</div>",
        unsafe_allow_html=True)

    for i, msg in hits:
        role_label = "You" if msg["role"] == "user" else "◆ NarrativeForge"
        content    = msg["content"]
        # Highlight match
        idx = content.lower().find(q)
        if idx >= 0:
            pre  = _html.escape(content[max(0, idx-60):idx])
            match = _html.escape(content[idx:idx+len(q)])
            post = _html.escape(content[idx+len(q):idx+len(q)+80])
            snippet = (f"…{pre}<mark style='background:rgba(77,107,254,0.25);"
                       f"color:var(--text-primary);border-radius:2px;'>{match}</mark>{post}…")
        else:
            snippet = _html.escape(content[:140])
        st.markdown(
            f"<div class='issue-card' style='margin-bottom:6px;'>"
            f"<div style='font-size:0.62rem;color:var(--primary);font-family:JetBrains Mono,monospace;"
            f"font-weight:600;margin-bottom:4px;'>MSG {i+1} · {role_label}</div>"
            f"<div style='font-size:0.78rem;color:var(--text-primary);line-height:1.55;'>"
            f"{snippet}</div></div>",
            unsafe_allow_html=True)


# ── First Draft Mode ──────────────────────────────────────────────────────────
_FIRST_DRAFT_JS = """
<script>
(function() {
  function lockInput() {
    const inputs = window.parent.document.querySelectorAll(
      '[data-testid="stChatInputTextArea"] textarea'
    );
    inputs.forEach(function(el) {
      if (el.dataset.fdLocked) return;
      el.dataset.fdLocked = "1";
      el.addEventListener('keydown', function(e) {
        if (e.key === 'Backspace' || e.key === 'Delete') {
          e.preventDefault();
          e.stopPropagation();
        }
      }, true);
    });
  }
  lockInput();
  setInterval(lockInput, 1500);
})();
</script>
"""

def _inject_first_draft_mode():
    import streamlit.components.v1 as _sc
    _sc.html(_FIRST_DRAFT_JS, height=0)


# ── Vocabulary Richness Stats ──────────────────────────────────────────────────
def _vocab_stats(story):
    """Returns dict of advanced vocabulary metrics."""
    all_msgs = story.get("messages", [])
    all_text = " ".join(m["content"] for m in all_msgs
                        if not m["content"].startswith("◆"))
    if not all_text.strip():
        return None
    words_raw  = re.findall(r"\b[a-z']{2,}\b", all_text.lower())
    total      = len(words_raw)
    unique     = len(set(words_raw))
    richness   = round(unique / total * 100, 1) if total else 0
    # Dialogue % — count quoted speech
    quoted = re.findall(r'["\u201c\u201d][^"\u201c\u201d]{4,}["\u201c\u201d]', all_text)
    dialogue_words = sum(len(q.split()) for q in quoted)
    dialogue_pct   = round(dialogue_words / total * 100, 1) if total else 0
    avg_word_len   = round(sum(len(w) for w in words_raw) / total, 1) if total else 0
    long_words     = sum(1 for w in words_raw if len(w) >= 7)
    long_pct       = round(long_words / total * 100, 1) if total else 0
    return {
        "total": total, "unique": unique, "richness": richness,
        "dialogue_pct": dialogue_pct, "avg_word_len": avg_word_len,
        "long_pct": long_pct,
    }


# ── AI Conflict & Stakes Analyzer ─────────────────────────────────────────────
def _conflict_prompt(story, characters):
    prose = _all_prose(story)[:800]
    char_names = ", ".join(c["name"] for c in characters) if characters else "unknown"
    return (
        f"Analyse this {story['genre']} story excerpt as a literary editor. "
        f"Characters: {char_names}.\n\nExcerpt:\n{prose or 'Story just beginning.'}\n\n"
        f"Respond in EXACTLY this format, nothing else:\n"
        f"CENTRAL CONFLICT: [one sentence — what is the core struggle?]\n"
        f"STAKES: [one sentence — what does the protagonist stand to lose?]\n"
        f"ANTAGONIST FORCE: [one sentence — person, system, nature, or inner demon?]\n"
        f"TENSION LEVEL: [1-10 score]\n"
        f"MISSING ELEMENT: [one sentence — what story element is weakest right now?]"
    )


def _sidebar_intelligence(story, characters, scenes, chapters=None, username=None):
    """
    The Intelligence tab — all 5 AI analysis tools in one panel.
    Tabs within tabs: Timeline | Show/Tell | Tension | Plot Holes | Voice | Tools
    """
    if chapters is None:
        chapters = []
    if not story.get("messages"):
        st.markdown(
            "<div style='color:#6B7080;font-size:0.82rem;text-align:center;padding:20px;'>"
            "Write some story first to unlock intelligence tools.</div>",
            unsafe_allow_html=True)
        return

    i1, i2, i3, i4, i5, i6, i7, i8, i9, i10, i11, i12, i13, i14, i15, i16 = st.tabs([
        "🗓 Timeline", "👁 Show/Tell", "⚡ Tension",
        "🕳 Plot Holes", "🎙 Voices", "🔧 Tools",
        "📈 Mood Arc", "🔍 Analyse", "⚔️ Conflict",
        "✏️ Grammar", "🎬 Screenplay", "📊 Analytics", "🤖 AI Suite",
        "🎭 Director", "🔀 CYOA", "🎨 Visuals",
    ])

    # ── Timeline ──────────────────────────────────────────────────────────────
    with i1:
        _render_timeline(chapters, scenes)

    # ── Show Don't Tell ───────────────────────────────────────────────────────
    with i2:
        st.markdown(
            "<div style='font-size:0.7rem;color:#6B7080;margin-bottom:8px;'>"
            "Scans your prose for telling instead of showing.</div>",
            unsafe_allow_html=True)
        sdt_key = f"sdt_result_{story['id']}"
        if st.button("👁 Run Show/Tell Scan", key="run_sdt", use_container_width=True, type="primary"):
            prose = _all_prose(story)
            st.session_state[sdt_key] = _show_dont_tell(prose)
        flagged = st.session_state.get(sdt_key)
        if flagged is None:
            st.caption("Click above to scan your prose.")
        elif not flagged:
            st.markdown(
                "<div style='color:#4D6BFE;font-size:0.82rem;text-align:center;"
                "padding:12px;'>✅ No tell-not-show patterns found!</div>",
                unsafe_allow_html=True)
        else:
            for f in flagged:
                st.markdown(f"""
                    <div class='issue-card' style='border-left-color:#fbbf24;margin-bottom:6px;'>
                        <div style='font-size:0.7rem;color:#fbbf24;font-weight:700;
                            margin-bottom:4px;'>⚠️ {_html.escape(f['issue'])}</div>
                        <div style='font-size:0.76rem;color:#A8BCFF;font-style:italic;
                            line-height:1.5;'>&ldquo;{_html.escape(f['sentence'][:100])}…&rdquo;</div>
                    </div>
                """, unsafe_allow_html=True)

    # ── Tension Meter ─────────────────────────────────────────────────────────
    with i3:
        ten_key = f"tension_result_{story['id']}"
        if st.button("⚡ Run Tension Analysis", key="run_tension", use_container_width=True, type="primary"):
            st.session_state[ten_key] = _tension_meter(story)
        tension = st.session_state.get(ten_key)
        if tension is None:
            st.caption("Click above to analyse tension.")
        elif not tension:
            st.caption("No AI-generated prose found.")
        else:
            st.markdown(
                "<div style='font-size:0.7rem;color:#6B7080;margin-bottom:10px;'>"
                "Higher = more tense. Aim for variation — peaks and valleys.</div>",
                unsafe_allow_html=True)
            max_score = max(t["score"] for t in tension) or 1
            for t in tension:
                score = t["score"]
                bar_w = int(score / max_score * 100)
                if   score >= 70: color, label = "#f87171", "tense"
                elif score >= 40: color, label = "#fbbf24", "building"
                else:             color, label = "#4D6BFE", "calm"
                st.markdown(f"""
                    <div style='display:flex;align-items:center;gap:8px;margin:5px 0;'>
                        <div style='width:52px;font-size:0.68rem;color:#6B7080;
                            font-family:monospace;'>{_html.escape(t['label'])}</div>
                        <div style='flex:1;background:rgba(77,107,254,0.06);
                            border-radius:4px;height:10px;'>
                            <div style='width:{bar_w}%;background:{color};height:10px;
                                border-radius:4px;transition:width 0.4s;'></div>
                        </div>
                        <div style='width:50px;font-size:0.68rem;color:{color};
                            text-align:right;font-weight:600;'>{score}% {label}</div>
                    </div>
                """, unsafe_allow_html=True)

    # ── Plot Hole Detector ────────────────────────────────────────────────────
    with i4:
        st.markdown(
            "<div style='font-size:0.7rem;color:#6B7080;margin-bottom:8px;'>"
            "AI reads your full story and flags narrative problems.</div>",
            unsafe_allow_html=True)
        if st.button("🕳 Run Plot Hole Analysis", use_container_width=True,
                     type="primary", key="run_plot_holes"):
            st.session_state.pop("plot_holes_result", None)
            with st.spinner("Reading your story arc…"):
                raw = _call_full(_plot_hole_prompt(story, characters, scenes),
                                 max_tokens=300)
            st.session_state["plot_holes_result"] = _parse_plot_holes(raw)
            st.rerun()

        ph = st.session_state.get("plot_holes_result")
        if ph:
            icons  = {"UNRESOLVED_THREADS": ("🧵", "#f87171"),
                      "CONTRADICTIONS":     ("⚡", "#fbbf24"),
                      "MISSING_MOTIVATION": ("❓", "#93c5fd"),
                      "TIMELINE_GAPS":      ("🕰", "#c4b5fd")}
            clean  = {"UNRESOLVED_THREADS": "Unresolved Threads",
                      "CONTRADICTIONS":     "Contradictions",
                      "MISSING_MOTIVATION": "Missing Motivation",
                      "TIMELINE_GAPS":      "Timeline Gaps"}
            any_found = False
            for key, (icon, color) in icons.items():
                val = ph.get(key, "").strip()
                if val and val.lower() != "none found.":
                    any_found = True
                    st.markdown(f"""
                        <div class='issue-card' style='border-left-color:{color};margin-bottom:6px;'>
                            <div style='font-size:0.72rem;color:{color};font-weight:700;
                                margin-bottom:4px;'>{icon} {clean[key]}</div>
                            <div style='font-size:0.78rem;color:var(--text-primary);line-height:1.5;'>
                                {_html.escape(val)}</div>
                        </div>
                    """, unsafe_allow_html=True)
            if not any_found:
                st.markdown(
                    "<div style='color:#4D6BFE;font-size:0.82rem;text-align:center;"
                    "padding:12px;'>✅ No plot holes detected!</div>",
                    unsafe_allow_html=True)

    # ── Character Voice Checker ───────────────────────────────────────────────
    with i5:
        if not characters:
            st.markdown(
                "<div style='color:#6B7080;font-size:0.82rem;text-align:center;"
                "padding:12px;'>Add characters in the Cast tab first.</div>",
                unsafe_allow_html=True)
        else:
            st.markdown(
                "<div style='font-size:0.7rem;color:#6B7080;margin-bottom:8px;'>"
                "AI checks whether each character speaks in a distinct voice.</div>",
                unsafe_allow_html=True)
            if st.button("🎙 Analyse Character Voices", use_container_width=True,
                         type="primary", key="run_voice_check"):
                st.session_state.pop("voice_check_result", None)
                prompt = _voice_check_prompt(story, characters)
                if prompt:
                    with st.spinner("Listening to your characters…"):
                        result = _call_full(prompt, max_tokens=250)
                    st.session_state["voice_check_result"] = result
                    st.rerun()

            vc = st.session_state.get("voice_check_result")
            if vc:
                lines = [l.strip() for l in vc.splitlines() if l.strip()]
                for line in lines:
                    if ":" in line and not line.startswith("OVERALL"):
                        name, rest = line.split(":", 1)
                        # Extract score
                        score_match = re.search(r'\b([1-9]|10)\b', rest)
                        score = int(score_match.group()) if score_match else 5
                        color = "#4D6BFE" if score >= 7 else "#fbbf24" if score >= 4 else "#f87171"
                        st.markdown(f"""
                            <div class='char-card' style='margin-bottom:5px;'>
                                <div style='display:flex;align-items:center;gap:8px;'>
                                    <div style='font-weight:700;color:var(--text-primary);
                                        font-size:0.85rem;flex:1;'>
                                        {_html.escape(name.strip())}</div>
                                    <div style='font-size:0.72rem;font-weight:800;
                                        color:{color};font-family:monospace;'>{score}/10</div>
                                </div>
                                <div style='font-size:0.76rem;color:#A8BCFF;
                                    line-height:1.5;margin-top:4px;font-style:italic;'>
                                    {_html.escape(rest.split("—", 1)[-1].strip() if "—" in rest else rest.strip())}
                                </div>
                            </div>
                        """, unsafe_allow_html=True)
                    elif line.startswith("OVERALL"):
                        st.markdown("---")
                        overall = line.split(":", 1)[-1].strip()
                        st.markdown(
                            f"<div style='font-size:0.8rem;color:#4D6BFE;font-style:italic;"
                            f"padding:6px 0;'>📊 {_html.escape(overall)}</div>",
                            unsafe_allow_html=True)

    # ── Tools (Title Gen, First Line, Genre Prompts) ──────────────────────────
    with i6:
        st.markdown(
            "<div style='font-size:0.7rem;color:#6B7080;margin-bottom:10px;'>"
            "Quick AI tools to get unstuck.</div>",
            unsafe_allow_html=True)

        # Title Generator
        st.markdown(
            "<div style='font-size:0.72rem;color:#6B7080;font-weight:700;"
            "text-transform:uppercase;letter-spacing:0.1em;margin-bottom:6px;'>"
            "📛 Title Generator</div>",
            unsafe_allow_html=True)
        if st.button("Generate 5 Titles", use_container_width=True, key="gen_titles"):
            prose_snip = _all_prose(story)[:600]
            prompt = (
                f"Generate exactly 5 compelling, literary titles for this {story['genre']} story "
                f"with a {story['tone']} tone. Each title should be on its own line, numbered. "
                f"No explanations. Story excerpt:\n{prose_snip or 'No content yet — base titles on genre and tone only.'}"
            )
            with st.spinner("Generating titles…"):
                result = _call_full(prompt, max_tokens=80)
            if result == llm.AI_UNAVAILABLE:
                _ai_error_msg()
            else:
                st.session_state["title_results"] = result
            st.rerun()
        if st.session_state.get("title_results"):
            st.markdown(
                "<div style='font-size:0.68rem;color:var(--text-muted);margin-bottom:6px;'>"
                "Click a title to apply it to your story:</div>",
                unsafe_allow_html=True)
            for line in st.session_state["title_results"].splitlines():
                line = line.strip()
                if not line:
                    continue
                display = re.sub(r"^\d+[\.\)]\s*", "", line).strip().strip('"').strip("'")
                if display and st.button(f"✦ {display}", key=f"title_apply_{hash(display) % 99999}",
                                         use_container_width=True):
                    _apply_title(story, display)
                    st.session_state["title_results"] = None
                    st.rerun()
        st.markdown("<div style='margin:10px 0;'></div>", unsafe_allow_html=True)

        # First Line Forge
        st.markdown(
            "<div style='font-size:0.72rem;color:var(--text-muted);font-weight:700;"
            "text-transform:uppercase;letter-spacing:0.1em;margin-bottom:6px;'>"
            "✍️ First Line Forge</div>",
            unsafe_allow_html=True)
        if st.button("Generate 5 Opening Lines", use_container_width=True, key="gen_firstlines"):
            char_names = [c["name"] for c in characters] if characters else []
            prompt = (
                f"Write exactly 5 distinct, powerful opening lines for a {story['genre']} story "
                f"with a {story['tone']} tone. "
                f"{'Characters: ' + ', '.join(char_names) + '. ' if char_names else ''}"
                "Each must be a single sentence on its own line, numbered. "
                "No explanations. Make each one unforgettable."
            )
            with st.spinner("Forging opening lines…"):
                result = _call_full(prompt, max_tokens=120)
            if result == llm.AI_UNAVAILABLE:
                _ai_error_msg()
            else:
                st.session_state["firstline_results"] = result
            st.rerun()
        if st.session_state.get("firstline_results"):
            st.markdown(
                f"<div style='background:var(--bg-card);border:1px solid rgba(77,107,254,0.18);"
                f"border-radius:8px;padding:12px 16px;font-size:0.85rem;color:var(--text-primary);"
                f"line-height:1.9;font-family:'Inter',sans-serif;'>"
                f"{_html.escape(st.session_state['firstline_results'])}</div>",
                unsafe_allow_html=True)
        st.markdown("<div style='margin:10px 0;'></div>", unsafe_allow_html=True)

        # Genre Prompt Generator
        st.markdown(
            "<div style='font-size:0.72rem;color:#6B7080;font-weight:700;"
            "text-transform:uppercase;letter-spacing:0.1em;margin-bottom:6px;'>"
            "💡 Writing Prompt (I'm Stuck)</div>",
            unsafe_allow_html=True)
        if st.button("Give Me 3 Prompts", use_container_width=True, key="gen_prompts"):
            prose_snip = _all_prose(story)[-400:] if story.get("messages") else ""
            char_names = [c["name"] for c in characters] if characters else []
            prompt = (
                f"Give exactly 3 specific, actionable writing prompts to continue "
                f"this {story['genre']} story (tone: {story['tone']}). "
                f"{'Characters: ' + ', '.join(char_names) + '. ' if char_names else ''}"
                f"{'Recent prose: ' + prose_snip if prose_snip else ''} "
                "Each prompt is 1-2 sentences, numbered, builds naturally from where the story is. "
                "No preamble."
            )
            with st.spinner("Thinking of prompts…"):
                result = _call_full(prompt, max_tokens=150)
            if result == llm.AI_UNAVAILABLE:
                _ai_error_msg()
            else:
                st.session_state["prompt_results"] = result
            st.rerun()
        if st.session_state.get("prompt_results"):
            st.markdown(
                f"<div style='background:var(--bg-card);border:1px solid rgba(77,107,254,0.18);"
                f"border-radius:8px;padding:12px 16px;font-size:0.85rem;color:var(--text-primary);"
                f"line-height:1.7;font-family:'Inter',sans-serif;'>"
                f"{_html.escape(st.session_state['prompt_results'])}</div>",
                unsafe_allow_html=True)

        st.markdown("<div style='margin:10px 0;'></div>", unsafe_allow_html=True)

        # ── Sentence Rewriter ─────────────────────────────────────────────────
        st.markdown(
            "<div style='font-size:0.72rem;color:#6B7080;font-weight:700;"
            "text-transform:uppercase;letter-spacing:0.1em;margin-bottom:6px;'>"
            "✂️ Sentence Rewriter</div>", unsafe_allow_html=True)
        st.markdown(
            "<div style='font-size:0.70rem;color:var(--text-muted);margin-bottom:8px;'>"
            "Paste any sentence for 3 fresh rewrites keeping the same meaning.</div>",
            unsafe_allow_html=True)
        rewrite_input = st.text_area("Sentence to rewrite", height=70,
                                     key="rewrite_input",
                                     label_visibility="collapsed",
                                     placeholder="Paste a sentence here…",
                                     max_chars=400)
        rw_style = st.selectbox(
            "Style", ["Same tone", "More lyrical", "Shorter & punchier",
                      "More tense", "Period-appropriate"],
            key="rewrite_style", label_visibility="collapsed")
        if st.button("✂️ Rewrite ×3", use_container_width=True, key="do_rewrite"):
            if rewrite_input.strip():
                prompt = (
                    f"Rewrite this sentence exactly 3 times. Style: {rw_style}. "
                    f"Keep the same meaning. Each version on its own numbered line. "
                    f"No preamble, no explanations.\n\nOriginal: {rewrite_input.strip()}"
                )
                with st.spinner("Rewriting…"):
                    result = _call_full(prompt, max_tokens=160)
                if result == llm.AI_UNAVAILABLE:
                    _ai_error_msg()
                else:
                    st.session_state["rewrite_results"] = result
                st.rerun()
            else:
                st.warning("Paste a sentence first.")
        rw = st.session_state.get("rewrite_results")
        if rw:
            for line in rw.splitlines():
                line = line.strip()
                if not line:
                    continue
                clean = re.sub(r"^\d+[\.\)]\s*", "", line).strip().strip('"').strip("'")
                if clean:
                    st.markdown(
                        f"<div class='issue-card' style='cursor:pointer;margin-bottom:5px;'>"
                        f"<div style='font-size:0.82rem;color:var(--text-primary);"
                        f"line-height:1.55;font-style:italic;'>{_html.escape(clean)}</div>"
                        f"</div>", unsafe_allow_html=True)
            if st.button("🗑 Clear", key="clear_rewrite", use_container_width=True):
                st.session_state.pop("rewrite_results", None)
                st.rerun()

    # ── Mood Arc (Sentiment) ──────────────────────────────────────────────────
    with i7:
        st.markdown(
            "<div style='font-size:0.7rem;color:var(--text-muted);margin-bottom:10px;'>"
            "Tracks the emotional tone of your story message by message — "
            "positive scores = hopeful/triumphant, negative = dark/tense.</div>",
            unsafe_allow_html=True)
        arc_key = f"arc_result_{story['id']}"
        if st.button("📈 Run Mood Arc", key="run_arc", use_container_width=True, type="primary"):
            st.session_state[arc_key] = _sentiment_arc(story)
        arc_data = st.session_state.get(arc_key)
        if arc_data is None:
            st.caption("Click above to generate mood arc.")
        elif not arc_data:
            st.markdown(
                "<div style='color:var(--text-muted);font-size:0.82rem;text-align:center;"
                "padding:20px;'>Write more story to see the mood arc.</div>",
                unsafe_allow_html=True)
        else:
            import pandas as pd
            df = pd.DataFrame(arc_data).set_index("Message")
            st.line_chart(df, color="#4D6BFE", height=220)
            scores   = [p["Sentiment"] for p in arc_data]
            avg_mood = round(sum(scores) / len(scores), 1)
            peak_msg = max(arc_data, key=lambda x: x["Sentiment"])
            dark_msg = min(arc_data, key=lambda x: x["Sentiment"])
            c1, c2, c3 = st.columns(3)
            c1.metric("Avg Mood", f"{avg_mood:+.1f}")
            c2.metric("Peak", f"Msg {peak_msg['Message']} ({peak_msg['Sentiment']:+.1f})")
            c3.metric("Darkest", f"Msg {dark_msg['Message']} ({dark_msg['Sentiment']:+.1f})")
            if avg_mood > 5:
                mood_label = "🌤 Generally uplifting"
            elif avg_mood < -5:
                mood_label = "🌩 Predominantly dark"
            else:
                mood_label = "⚖️ Balanced tension"
            st.markdown(
                f"<div style='font-size:0.8rem;color:var(--primary);margin-top:6px;'>"
                f"{mood_label}</div>", unsafe_allow_html=True)

    # ── Analyse: Cliché Detector + Reading Level ──────────────────────────────
    with i8:
        all_text = " ".join(m["content"] for m in story.get("messages", []))
        rl_key = f"rl_result_{story['id']}_{len(all_text)}"

        # Reading Level
        st.markdown(
            "<div style='font-size:0.72rem;color:var(--text-muted);font-weight:700;"
            "text-transform:uppercase;letter-spacing:0.1em;margin-bottom:8px;'>"
            "📚 Reading Level</div>", unsafe_allow_html=True)
        if len(all_text.split()) >= 30:
            if rl_key not in st.session_state:
                st.session_state[rl_key] = _reading_level(all_text)
            rl = st.session_state[rl_key]
            if rl:
                r1, r2, r3 = st.columns(3)
                r1.metric("Ease Score", f"{rl['ease']}/100")
                r2.metric("Grade Level", f"Grade {rl['grade']}")
                r3.metric("Sentences", rl["sentences"])
                st.markdown(
                    f"<div class='insight-card'>"
                    f"<div class='insight-title'>Audience</div>"
                    f"<div class='insight-body'>{rl['label']}</div>"
                    f"</div>", unsafe_allow_html=True)
                if rl["ease"] < 50:
                    st.info("💡 Try shorter sentences to improve readability.")
                elif rl["ease"] > 80:
                    st.info("💡 Your prose is very accessible — consider adding complexity for literary depth.")
        else:
            st.markdown(
                "<div style='color:var(--text-muted);font-size:0.82rem;'>Write at least 30 words to analyse reading level.</div>",
                unsafe_allow_html=True)

        st.markdown("<div style='margin:14px 0;height:1px;background:var(--border-subtle);'></div>",
                    unsafe_allow_html=True)

        # Cliché Detector
        st.markdown(
            "<div style='font-size:0.72rem;color:var(--text-muted);font-weight:700;"
            "text-transform:uppercase;letter-spacing:0.1em;margin-bottom:8px;'>"
            "🚫 Cliché Detector</div>", unsafe_allow_html=True)
        if st.button("Scan for Clichés", use_container_width=True, key="scan_cliches"):
            found = _detect_cliches(all_text)
            st.session_state[f"cliche_results_{story['id']}"] = found
        results = st.session_state.get(f"cliche_results_{story['id']}")
        if results is not None:
            if not results:
                st.markdown(
                    "<div style='color:var(--success);font-size:0.85rem;text-align:center;"
                    "padding:12px;'>✅ No common clichés found! Fresh writing.</div>",
                    unsafe_allow_html=True)
            else:
                st.markdown(
                    f"<div style='color:var(--danger);font-size:0.80rem;margin-bottom:8px;'>"
                    f"Found {len(results)} cliché(s):</div>", unsafe_allow_html=True)
                for c in results:
                    st.markdown(
                        f"<div class='issue-card' style='border-left-color:var(--warning);'>"
                        f"<div style='font-size:0.78rem;color:var(--warning);'>⚠ \"{_html.escape(c)}\"</div>"
                        f"<div style='font-size:0.72rem;color:var(--text-muted);margin-top:3px;'>"
                        f"Consider rewriting this for originality.</div>"
                        f"</div>", unsafe_allow_html=True)

        st.markdown("<div style='margin:14px 0;height:1px;background:var(--border-subtle);'></div>",
                    unsafe_allow_html=True)

        # Word frequency
        st.markdown(
            "<div style='font-size:0.72rem;color:var(--text-muted);font-weight:700;"
            "text-transform:uppercase;letter-spacing:0.1em;margin-bottom:8px;'>"
            "📊 Top Words</div>", unsafe_allow_html=True)
        words_raw = re.findall(r"\b[a-z]{4,}\b", all_text.lower())
        common    = Counter(w for w in words_raw if w not in _STOPWORDS).most_common(10)
        if common:
            max_c = common[0][1]
            for word, cnt in common:
                bar_w = int((cnt / max_c) * 100)
                st.markdown(
                    f"<div style='display:flex;align-items:center;gap:8px;margin-bottom:5px;'>"
                    f"<div style='width:80px;font-size:0.75rem;color:var(--text-primary);'>{word}</div>"
                    f"<div style='flex:1;background:var(--primary-dim);border-radius:4px;height:8px;'>"
                    f"<div style='background:var(--primary);width:{bar_w}%;height:8px;border-radius:4px;'></div></div>"
                    f"<div style='width:24px;font-size:0.70rem;color:var(--text-muted);'>{cnt}</div>"
                    f"</div>", unsafe_allow_html=True)
        else:
            st.markdown("<div style='color:var(--text-muted);font-size:0.8rem;'>Write more to see word frequency.</div>",
                        unsafe_allow_html=True)

    # ── Conflict & Stakes Analyser ────────────────────────────────────────────
    with i9:
        st.markdown(
            "<div style='font-size:0.7rem;color:var(--text-muted);margin-bottom:10px;'>"
            "AI maps your story's central conflict, stakes, and what's missing.</div>",
            unsafe_allow_html=True)
        if st.button("⚔️ Analyse Conflict & Stakes", use_container_width=True,
                     type="primary", key="run_conflict"):
            st.session_state.pop("conflict_result", None)
            with st.spinner("Mapping your story conflict…"):
                raw = _call_full(_conflict_prompt(story, characters), max_tokens=200)
            if raw == llm.AI_UNAVAILABLE:
                _ai_error_msg()
            else:
                st.session_state["conflict_result"] = raw
            st.rerun()

        cr = st.session_state.get("conflict_result")
        if cr:
            labels = {
                "CENTRAL CONFLICT":  ("⚔️", "#f87171"),
                "STAKES":            ("💀", "#fbbf24"),
                "ANTAGONIST FORCE":  ("🎭", "#c4b5fd"),
                "TENSION LEVEL":     ("📊", "#4D6BFE"),
                "MISSING ELEMENT":   ("🔍", "#93c5fd"),
            }
            parsed = {}
            for line in cr.splitlines():
                for key in labels:
                    if line.upper().startswith(key + ":"):
                        parsed[key] = line.split(":", 1)[1].strip()
                        break

            if not parsed:
                st.markdown(
                    f"<div class='issue-card'><div style='font-size:0.82rem;"
                    f"color:var(--text-primary);line-height:1.6;'>{_html.escape(cr)}</div></div>",
                    unsafe_allow_html=True)
            else:
                for key, (icon, color) in labels.items():
                    val = parsed.get(key, "")
                    if not val:
                        continue
                    # Special rendering for tension level — show bar
                    if key == "TENSION LEVEL":
                        score_match = re.search(r'\b([1-9]|10)\b', val)
                        score = int(score_match.group()) if score_match else 5
                        bar_color = "#f87171" if score >= 8 else "#fbbf24" if score >= 5 else "#4D6BFE"
                        st.markdown(
                            f"<div class='issue-card' style='border-left-color:{bar_color};margin-bottom:6px;'>"
                            f"<div style='font-size:0.72rem;color:{bar_color};font-weight:700;"
                            f"margin-bottom:6px;'>{icon} Tension Level</div>"
                            f"<div style='display:flex;align-items:center;gap:8px;'>"
                            f"<div style='flex:1;background:var(--primary-dim);border-radius:4px;height:10px;'>"
                            f"<div style='width:{score*10}%;background:{bar_color};height:10px;border-radius:4px;'></div>"
                            f"</div><div style='font-size:0.85rem;font-weight:800;color:{bar_color};"
                            f"font-family:JetBrains Mono,monospace;min-width:32px;'>{score}/10</div>"
                            f"</div></div>", unsafe_allow_html=True)
                    else:
                        st.markdown(
                            f"<div class='issue-card' style='border-left-color:{color};margin-bottom:6px;'>"
                            f"<div style='font-size:0.72rem;color:{color};font-weight:700;"
                            f"margin-bottom:4px;'>{icon} {key.title()}</div>"
                            f"<div style='font-size:0.80rem;color:var(--text-primary);line-height:1.55;'>"
                            f"{_html.escape(val)}</div></div>",
                            unsafe_allow_html=True)


    # ── Grammar & Style Checker (i10) ─────────────────────────────────────────
    with i10:
        if _HAS_GRAMMAR:
            if _HAS_ACHIEVEMENTS:
                record_analysis(st.session_state.get("username", ""))
            show_grammar_checker(story, ctx="intel")
        else:
            st.info("grammar.py module not found.")

    # ── Screenplay Converter (i11) ────────────────────────────────────────────
    with i11:
        if _HAS_SCREENPLAY:
            show_screenplay_tab(story, characters, ctx="intel")
        else:
            st.info("screenplay.py module not found.")

    # ── Analytics Dashboard (i12) ─────────────────────────────────────────────
    with i12:
        if _HAS_ANALYTICS:
            show_full_analytics(story, characters, chapters, scenes)
        else:
            st.info("analytics.py module not found.")

    # ── AI Creative Suite (i13) ───────────────────────────────────────────────
    with i13:
        if _HAS_AI_TOOLS:
            show_ai_tools_panel(story, characters, ctx="intel")
        else:
            st.info("ai_tools.py module not found.")

    # ── Director's Cut (i14) ──────────────────────────────────────────────────
    with i14:
        if _HAS_DIRECTORS_CUT:
            show_directors_cut(story, ctx="intel")
        else:
            st.info("directors_cut.py module not found.")

    # ── Interactive / CYOA (i15) ──────────────────────────────────────────────
    with i15:
        if _HAS_INTERACTIVE:
            show_interactive_story(story, characters, username, ctx="intel")
        else:
            st.info("interactive_story.py module not found.")

    # ── Visual Generator (i16) ────────────────────────────────────────────────
    with i16:
        if _HAS_VISUAL:
            show_visual_generator(story, characters, ctx="intel")
        else:
            st.info("visual_gen.py module not found.")



# ── Main workspace ─────────────────────────────────────────────────────────────

def _inject_autosave(story_id):
    """
    Auto-save chat input to localStorage every 30s.
    Shows a recovery banner if unsaved draft is found on load.
    """
    import streamlit.components.v1 as _sc
    _sc.html(f"""
        <script>
        (function() {{
          var DRAFT_KEY = 'nf_draft_{story_id}';
          var DRAFT_TS  = 'nf_draft_ts_{story_id}';

          // ── On load: check for unsaved draft ────────────────────────────
          var saved = localStorage.getItem(DRAFT_KEY);
          var ts    = localStorage.getItem(DRAFT_TS);
          if (saved && saved.length > 5) {{
            var ageMin = ts ? Math.round((Date.now() - parseInt(ts)) / 60000) : '?';
            var banner = document.createElement('div');
            banner.id  = 'nf-draft-banner';
            banner.style.cssText = 'position:fixed;bottom:80px;right:20px;z-index:8888;'
              + 'background:#161820;border:1px solid rgba(77,107,254,0.5);'
              + 'border-radius:10px;padding:12px 16px;font-family:Inter,sans-serif;'
              + 'font-size:0.82rem;color:#C5C8D4;box-shadow:0 4px 20px rgba(0,0,0,0.5);'
              + 'max-width:280px;';
            banner.innerHTML = '<div style="color:#A8BCFF;font-weight:700;margin-bottom:6px;">📝 Unsaved Draft Found</div>'
              + '<div style="color:#6B7080;font-size:0.75rem;margin-bottom:10px;">'
              + 'Saved ' + ageMin + 'm ago: ' + saved.substring(0,60) + '…</div>'
              + '<div style="display:flex;gap:8px;">'
              + '<button id="nf-draft-restore" style="background:#4D6BFE;color:#fff;border:none;'
              + 'border-radius:6px;padding:5px 12px;cursor:pointer;font-size:0.78rem;">Restore</button>'
              + '<button id="nf-draft-discard" style="background:transparent;color:#6B7080;border:1px solid #333;'
              + 'border-radius:6px;padding:5px 12px;cursor:pointer;font-size:0.78rem;">Discard</button>'
              + '</div>';
            document.body.appendChild(banner);

            document.getElementById('nf-draft-restore').onclick = function() {{
              var ta = window.parent.document.querySelector(
                'textarea[data-testid="stChatInputTextArea"], .stChatInput textarea');
              if (ta) {{
                var nativeInputSetter = Object.getOwnPropertyDescriptor(
                  window.parent.HTMLTextAreaElement.prototype, 'value').set;
                nativeInputSetter.call(ta, saved);
                ta.dispatchEvent(new Event('input', {{ bubbles: true }}));
                ta.focus();
              }}
              banner.remove();
            }};
            document.getElementById('nf-draft-discard').onclick = function() {{
              localStorage.removeItem(DRAFT_KEY);
              localStorage.removeItem(DRAFT_TS);
              banner.remove();
            }};
          }}

          // ── Auto-save every 30s ──────────────────────────────────────────
          function saveDraft() {{
            var ta = window.parent.document.querySelector(
              'textarea[data-testid="stChatInputTextArea"], .stChatInput textarea');
            if (ta && ta.value && ta.value.length > 3) {{
              localStorage.setItem(DRAFT_KEY, ta.value);
              localStorage.setItem(DRAFT_TS, Date.now().toString());
            }}
          }}

          // Clear draft when message is sent
          function clearDraftOnSend() {{
            var btn = window.parent.document.querySelector(
              'button[data-testid="stChatInputSubmitButton"]');
            if (btn) {{
              btn.addEventListener('click', function() {{
                localStorage.removeItem(DRAFT_KEY);
                localStorage.removeItem(DRAFT_TS);
              }}, {{ once: false }});
            }}
          }}

          setInterval(saveDraft, 30000);
          setTimeout(clearDraftOnSend, 2000);
        }})();
        </script>
    """, height=0)

def _inject_keyboard_shortcuts():
    """Full keyboard shortcut suite with cheat sheet overlay."""
    import streamlit.components.v1 as _sc
    _sc.html("""
        <script>
        (function() {
          if (window._nf_shortcuts_loaded) return;
          window._nf_shortcuts_loaded = true;

          function clickBtn(testid) {
            var btn = window.parent.document.querySelector(
              '[data-testid="' + testid + '"] button, button[key="' + testid + '"]');
            if (btn) btn.click();
          }
          function focusInput() {
            var ta = window.parent.document.querySelector(
              'textarea[data-testid="stChatInputTextArea"], .stChatInput textarea');
            if (ta) { ta.focus(); ta.scrollIntoView({behavior:'smooth'}); }
          }

          document.addEventListener('keydown', function(e) {
            var tag = document.activeElement ? document.activeElement.tagName : '';
            var isInput = (tag === 'INPUT' || tag === 'TEXTAREA' || tag === 'SELECT'
                           || document.activeElement.isContentEditable);

            // Shift+? — cheat sheet (always works)
            if (e.shiftKey && e.key === '?') {
              e.preventDefault();
              var ex = document.getElementById('nf-shortcuts-overlay');
              if (ex) { ex.remove(); return; }
              var ov = document.createElement('div');
              ov.id = 'nf-shortcuts-overlay';
              ov.style.cssText = 'position:fixed;inset:0;background:rgba(0,0,0,0.6);z-index:9998;';
              ov.innerHTML = `
                <div style="position:absolute;top:50%;left:50%;transform:translate(-50%,-50%);
                  background:#161820;border:1px solid rgba(77,107,254,0.5);border-radius:16px;
                  padding:28px 36px;min-width:380px;max-width:90vw;
                  box-shadow:0 12px 60px rgba(0,0,0,0.7);font-family:'Inter',sans-serif;">
                  <div style="font-size:1.1rem;font-weight:700;color:#A8BCFF;margin-bottom:20px;">
                    ⌨️ Keyboard Shortcuts</div>
                  <table style="width:100%;border-collapse:collapse;font-size:0.84rem;color:#C5C8D4;">
                    <tr style="border-bottom:1px solid rgba(77,107,254,0.1);">
                      <th style="text-align:left;padding:4px 0;color:#6B7080;font-size:0.7rem;
                        text-transform:uppercase;letter-spacing:0.1em;">Navigation</th><th></th></tr>
                    <tr><td style="padding:5px 16px 5px 0;color:#4D6BFE;font-family:monospace;
                      font-weight:700;">Shift+?</td><td>Show/hide this cheat sheet</td></tr>
                    <tr><td style="padding:5px 16px 5px 0;color:#4D6BFE;font-family:monospace;
                      font-weight:700;">Esc</td><td>Close overlay</td></tr>
                    <tr style="border-bottom:1px solid rgba(77,107,254,0.1);">
                      <th style="text-align:left;padding:8px 0 4px;color:#6B7080;font-size:0.7rem;
                        text-transform:uppercase;letter-spacing:0.1em;">Writing</th><th></th></tr>
                    <tr><td style="padding:5px 16px 5px 0;color:#4D6BFE;font-family:monospace;
                      font-weight:700;">Alt+F</td><td>Focus chat input</td></tr>
                    <tr><td style="padding:5px 16px 5px 0;color:#4D6BFE;font-family:monospace;
                      font-weight:700;">Alt+1</td><td>Mode: Continue</td></tr>
                    <tr><td style="padding:5px 16px 5px 0;color:#4D6BFE;font-family:monospace;
                      font-weight:700;">Alt+2</td><td>Mode: Paragraph</td></tr>
                    <tr><td style="padding:5px 16px 5px 0;color:#4D6BFE;font-family:monospace;
                      font-weight:700;">Alt+3</td><td>Mode: Dialogue</td></tr>
                    <tr><td style="padding:5px 16px 5px 0;color:#4D6BFE;font-family:monospace;
                      font-weight:700;">Alt+4</td><td>Mode: Twist</td></tr>
                    <tr><td style="padding:5px 16px 5px 0;color:#4D6BFE;font-family:monospace;
                      font-weight:700;">Alt+5</td><td>Mode: Monologue</td></tr>
                    <tr style="border-bottom:1px solid rgba(77,107,254,0.1);">
                      <th style="text-align:left;padding:8px 0 4px;color:#6B7080;font-size:0.7rem;
                        text-transform:uppercase;letter-spacing:0.1em;">Features</th><th></th></tr>
                    <tr><td style="padding:5px 16px 5px 0;color:#4D6BFE;font-family:monospace;
                      font-weight:700;">Alt+R</td><td>Toggle Reading Mode</td></tr>
                    <tr><td style="padding:5px 16px 5px 0;color:#4D6BFE;font-family:monospace;
                      font-weight:700;">Alt+S</td><td>Save Snapshot (quick)</td></tr>
                    <tr><td style="padding:5px 16px 5px 0;color:#4D6BFE;font-family:monospace;
                      font-weight:700;">Alt+D</td><td>Back to Dashboard</td></tr>
                  </table>
                  <div style="margin-top:18px;font-size:0.72rem;color:#6B7080;text-align:center;">
                    Click anywhere · Press Esc · Press Shift+? to close</div>
                </div>`;
              ov.onclick = function(ev) {
                if (ev.target === ov) ov.remove();
              };
              document.body.appendChild(ov);
              return;
            }

            // Escape — close overlay
            if (e.key === 'Escape') {
              var ov = document.getElementById('nf-shortcuts-overlay');
              if (ov) { ov.remove(); return; }
            }

            if (isInput) return; // rest only fire when NOT typing

            // Alt+F — focus chat input
            if (e.altKey && e.key === 'f') {
              e.preventDefault(); focusInput(); return;
            }
            // Alt+R — reading mode toggle
            if (e.altKey && e.key === 'r') {
              e.preventDefault();
              var readBtn = window.parent.document.querySelector('[data-testid="topnav_read"]');
              if (readBtn) readBtn.click();
              return;
            }
            // Alt+D — dashboard
            if (e.altKey && e.key === 'd') {
              e.preventDefault();
              var backBtn = window.parent.document.querySelector('[data-testid="topnav_back"]');
              if (backBtn) backBtn.click();
              return;
            }
            // Alt+1-5 mode buttons
            var modeMap = {'1':'mode_continue','2':'mode_paragraph',
                           '3':'mode_dialogue','4':'mode_twist','5':'mode_monologue'};
            if (e.altKey && modeMap[e.key]) {
              e.preventDefault();
              var btn = window.parent.document.querySelector(
                'button[data-testid="' + modeMap[e.key] + '"]');
              if (btn) btn.click();
              return;
            }
          });
        })();
        </script>
    """, height=0)


def _sidebar_settings_exports_fallback(story, characters, scenes):
    """Fallback export section when export.py is unavailable."""
    st.caption("export.py module not found — using built-in exports.")


def _force_sidebar_open():
    """Inject lightweight JS to open sidebar — zero blocking height."""
    st.markdown("""
        <script>
        setTimeout(function(){
            var btn=window.parent.document.querySelector('[data-testid="collapsedControl"] button');
            if(btn){btn.click();}
        },100);
        </script>
    """, unsafe_allow_html=True)

def show_workspace():
    workspace_style()
    # Only inject sidebar JS once per story session — not on every rerun
    _sidebar_opened_key = f"_sidebar_opened_{st.session_state.get('current_story','')}"
    if not st.session_state.get(_sidebar_opened_key):
        _force_sidebar_open()
        st.session_state[_sidebar_opened_key] = True
    story = _get_story()
    if not story:
        st.error("Story not found.")
        if st.button("← Back"):
            st.session_state.current_view = "dashboard"; st.rerun()
        return

    username   = st.session_state.username

    # ── Book mode — typeset book view with edit + PDF export ─────────────────
    if st.session_state.get("book_mode") and _HAS_BOOK_FEATURES:
        show_book_mode(story, username, save_story, characters=_cached_characters(username, story["id"]))
        return

    # ── Reading mode — full screen clean prose view ────────────────────────────
    if st.session_state.get("reading_mode"):
        show_reading_mode()
        return

    # ── Session tracker ────────────────────────────────────────────────────────
    sess_key   = f"_sess_start_{story['id']}"
    words_key  = f"_sess_words_{story['id']}"
    if sess_key not in st.session_state:
        st.session_state[sess_key]  = time.time()
        st.session_state[words_key] = _word_count(story)
    elapsed_min  = max(0, (time.time() - st.session_state[sess_key]) / 60)
    words_now    = _word_count(story)
    words_added  = max(0, words_now - st.session_state[words_key])
    wpm          = round(words_added / elapsed_min) if elapsed_min > 0.1 else 0

    # Track words for analytics/streaks (record only when words added)
    if words_added > 0 and _HAS_ANALYTICS:
        record_session_words(story["id"], words_added)
        st.session_state[words_key] = words_now  # update baseline

    # ── Load story data ONCE — reused by all sidebar tabs + nav bar ───────────
    _all_chars  = _cached_characters(username, story["id"])
    _all_scenes = _cached_scenes(username, story["id"])
    _all_chaps  = _cached_chapters(username, story["id"])

    # ── Top navigation bar ────────────────────────────────────────────────────
    nav1, nav2, nav3, nav4 = st.columns([2, 3, 2, 3])
    with nav1:
        if st.button("← Dashboard", key="topnav_back", use_container_width=True):
            save_story(username, story)
            st.session_state.current_view = "dashboard"
            st.session_state.show_cowrite_options = False
            _bust_cache()
            st.rerun()
    with nav2:
        bm_active = st.session_state.get("book_mode", False)
        bm_label  = "📖 Book Mode ✓" if bm_active else "📖 Book Mode"
        if st.button(bm_label, key="topnav_book", use_container_width=True):
            st.session_state["book_mode"] = not bm_active
            st.session_state.pop("reading_mode", None)
            _bust_cache()
            st.rerun()
    with nav3:
        bible_cache_key = f"_bible_cache_{story['id']}"
        bible_wc_key    = f"_bible_wc_{story['id']}"
        current_wc      = _word_count(story)
        bible_bytes     = st.session_state.get(bible_cache_key)
        # Only rebuild when word count changes AND bible was already generated
        if bible_bytes and st.session_state.get(bible_wc_key) != current_wc:
            bible_buf = _export_story_bible(story, _all_chars, _all_scenes, username)
            bible_bytes = bible_buf.getvalue()
            st.session_state[bible_cache_key] = bible_bytes
            st.session_state[bible_wc_key]    = current_wc
        if bible_bytes:
            st.download_button(
                "📚 Story Bible",
                data=bible_bytes,
                file_name=f"{story['title'].replace(' ','_')}_Bible.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True, key="dl_bible")
        else:
            if st.button("📚 Story Bible", use_container_width=True, key="gen_bible"):
                with st.spinner("Building…"):
                    bible_buf = _export_story_bible(story, _all_chars, _all_scenes, username)
                    st.session_state[bible_cache_key] = bible_buf.getvalue()
                    st.session_state[bible_wc_key]    = current_wc
                st.rerun()
    with nav4:
        st.markdown(f"""
            <div style='display:flex;gap:10px;justify-content:flex-end;padding-top:4px;flex-wrap:wrap;'>
                <span style='font-size:0.72rem;color:#6B7080;'>⏱ {int(elapsed_min)}m session</span>
                <span style='font-size:0.72rem;color:var(--primary);font-weight:600;'>+{words_added} words</span>
                {'<span style="font-size:0.72rem;color:var(--accent);">🚀 ' + str(wpm) + ' wpm</span>' if wpm > 0 else ''}
            </div>
        """, unsafe_allow_html=True)

    # ── Sidebar ────────────────────────────────────────────────────────────────
    with st.sidebar:
        st.markdown(f"### 👤 {username}")
        st.markdown("---")
        t1, t2, t3, t4, t5, t6, t7, t8, t9, t10, t11, t12, t13, t14, t15, t16, t17, t18, t19, t20, t21, t22 = st.tabs([
            "⚙️ Set", "👥 Cast", "🎬 Scenes", "📖 Chaps",
            "📈 Arc", "💊 Health", "📓 Notes", "📸 Snaps",
            "🔗 Map", "🧠 Intel", "🎭 Chat", "🔍 Find",
            "📥 Export", "🏅 Awards", "📅 Track", "🔍 Style",
            "🕰 Versions", "🧬 Chars+", "🌍 World", "📁 Import", "👥 Beta", "🎨 Visuals"])
        with t1: _sidebar_settings(story)
        with t2: _sidebar_characters(story)
        with t3: _sidebar_scenes(story)
        with t4: _sidebar_chapters(story)
        with t5: _sidebar_arc(story)
        with t6:
            _sidebar_health(story, _all_chars)
        with t7: _sidebar_notes(story)
        with t8: _sidebar_snapshots(story)
        with t9:
            if _all_chars:
                st.markdown(
                    "<div style='font-size:0.7rem;color:#6B7080;margin-bottom:6px;"
                    "text-align:center;'>Drag nodes · Hover for details · Scroll to zoom</div>",
                    unsafe_allow_html=True)
                render_relationship_map(_all_chars, _all_scenes, height=450)
            else:
                st.markdown(
                    "<div style='color:#6B7080;font-size:0.82rem;text-align:center;"
                    "padding:24px;'>Add characters in the Cast tab to see the relationship map.</div>",
                    unsafe_allow_html=True)
        with t10:
            _sidebar_intelligence(story, _all_chars, _all_scenes, _all_chaps, username)
        with t11:
            _sidebar_interview(story, _all_chars)
        with t12:
            _sidebar_search_story(story)
        with t13:
            if _HAS_EXPORT:
                show_export_panel(story, _all_chars, _all_scenes, username)
                if _HAS_ACHIEVEMENTS:
                    record_export(username)
            else:
                # Fallback export using existing functions
                st.markdown("**📥 Export Story**")
                _sidebar_settings_exports_fallback(story, _all_chars, _all_scenes)
        with t14:
            if _HAS_ACHIEVEMENTS:
                all_stories = st.session_state.get("stories", [])
                # Check for new achievements
                new_achs = check_and_award(all_stories, username)
                for ach in new_achs:
                    show_new_achievement(ach)
                show_achievements_panel(all_stories, username)
            else:
                st.info("achievements.py module not found.")
        with t15:
            if _HAS_TRACKER:
                show_daily_tracker(username)
            else:
                st.info("daily_tracker.py module not found.")
        with t16:
            if _HAS_GRAMMAR:
                show_grammar_checker(story, ctx="sidebar")
            else:
                st.info("grammar.py module not found.")
        with t17:
            if _HAS_VERSION_TIMELINE:
                show_version_timeline(story, username)
            else:
                st.info("version_timeline.py module not found.")
        with t18:
            if _HAS_CHAR_SUITE:
                show_character_suite(story, _all_chars)
            else:
                st.info("character_suite.py module not found.")
        with t19:
            if _HAS_WORLD_BUILDER:
                show_world_builder(story, username)
            else:
                st.info("world_builder.py module not found.")
        with t20:
            if _HAS_IMPORTER:
                show_scrivener_import(story, username)
            else:
                st.info("scrivener_import.py module not found.")
        with t21:
            if _HAS_BETA:
                show_beta_panel(story, username)
            else:
                st.info("beta_reading.py module not found.")
        with t22:
            if _HAS_VISUAL:
                show_visual_generator(story, _all_chars, ctx="sidebar")
            else:
                st.info("visual_gen.py module not found.")

    # ── Inject keyboard shortcuts (power user) ───────────────────────────────
    _inject_keyboard_shortcuts()
    # ── Auto-save chat input drafts to localStorage ──────────────────────────
    _inject_autosave(story["id"])

    # ── Reuse pre-loaded data for header, prompt, and export ────────────────
    characters = _all_chars
    scenes     = _all_scenes

    # ── Edit Mode toggle ──────────────────────────────────────────────────────
    edit_col, _ = st.columns([2, 8])
    with edit_col:
        edit_active = st.session_state.get("_ws_edit_mode", False)
        edit_label  = "✅ Exit Edit Mode" if edit_active else "✏️ Edit Messages"
        if st.button(edit_label, key="ws_edit_toggle", use_container_width=True):
            st.session_state["_ws_edit_mode"] = not edit_active
            st.rerun()

    # ── AI status badge (skip Ollama check when Groq is active) ─────────────
    if not llm.GROQ_API_KEY.strip():
        from utils import ollama_status_banner
        _ollama_ready = ollama_status_banner(OLLAMA_MODEL)
    else:
        st.markdown(
            f"<div style='display:inline-flex;align-items:center;gap:6px;"
            f"background:rgba(77,107,254,0.07);border:1px solid rgba(77,107,254,0.20);"
            f"border-radius:20px;padding:3px 11px;font-size:0.68rem;font-weight:700;"
            f"color:#4D6BFE;letter-spacing:0.08em;font-family:JetBrains Mono,monospace;'>"
            f"<span style='width:6px;height:6px;background:#4ADE80;border-radius:50%;"
            f"box-shadow:0 0 5px #4ADE80;'></span> ⚡ Groq · ready</div>",
            unsafe_allow_html=True)
        _ollama_ready = True

    # ── Header ─────────────────────────────────────────────────────────────────
    arc  = story.get("plot_arc", {})
    done = sum(1 for k,_ in PLOT_STAGES if arc.get(k))
    # FIX #2: escape all user-controlled values before embedding in HTML
    voice_badge = f"<span class='ws-badge'>✍️ {_html.escape(story['writing_style'][:22])}…</span>" if story.get("writing_style") else ""
    arc_badge   = f"<span class='ws-badge ws-badge-green'>📈 Arc {done}/5</span>" if done else ""
    st.markdown(f"""
        <div style='margin-bottom:6px;'>
            <span style='font-family:'Inter',sans-serif;font-size:1.9rem;font-weight:700;color:#4D6BFE;'>
                📖 {_html.escape(story['title'])}</span>
        </div>
        <div style='display:flex;gap:8px;margin-bottom:16px;flex-wrap:wrap;'>
            <span class='ws-badge'>📚 {_html.escape(story['genre'])}</span>
            <span class='ws-badge'>🎭 {_html.escape(story['tone'])}</span>
            <span class='ws-badge'>🤖 {_html.escape(llm.get_provider_label())}</span>
            {arc_badge}{voice_badge}
        </div>
    """, unsafe_allow_html=True)

    # ── Edit mode — show editor instead of normal chat ──────────────────────
    if st.session_state.get("_ws_edit_mode") and _HAS_BOOK_FEATURES:
        show_message_editor(story, username, save_story)
        st.stop()

    # ── Chat history ─────────────────────────────────────────────────────────
    if story.get("messages"):
        msgs = story["messages"]

        # Render all messages except the very last AI one individually
        non_last_parts = []
        for i, msg in enumerate(msgs):
            is_last = (i == len(msgs) - 1)
            if msg["role"] == "user":
                non_last_parts.append(
                    f"<div class='bubble-user'><div class='lbl'>🧑\u200d💻 You</div>"
                    f"<div class='txt'>{_html.escape(msg['content'])}</div></div>"
                )
            else:
                msg_content = msg["content"]
                is_redirect = (
                    msg_content.startswith("I'm here to help") or
                    msg_content.startswith("I'm designed to") or
                    msg_content.startswith("I focus only") or
                    msg_content.startswith("I can't help with that") or
                    msg_content.startswith("That's outside") or
                    (msg_content.startswith("◆") and len(msg_content) < 400)
                )
                safe_c = _html.escape(msg_content)
                if is_redirect:
                    non_last_parts.append(f"<div class='redirect-msg'>{safe_c}</div>")
                else:
                    non_last_parts.append(
                        f"<div class='bubble-ai'><div class='lbl'>◆ NarrativeForge</div>"
                        f"<div class='txt'>{safe_c}</div></div>"
                    )
            if not is_last:
                continue

            # ── Flush all messages so far ──────────────────────────────────
            if non_last_parts:
                st.markdown("".join(non_last_parts), unsafe_allow_html=True)
                non_last_parts = []

            # ── If last message is AI — show action buttons below it ───────
            if msg["role"] == "assistant":
                b1, b2 = st.columns(2)
                last_content = msg["content"]
                with b1:
                    if st.button("🔄 Regenerate", key=f"regen_{i}", use_container_width=True):
                        prev   = story["messages"][i-1]["content"] if i > 0 else ""
                        prompt = _build_prompt(prev, story, "paragraph", characters, scenes)
                        slot   = st.empty()
                        tokens = []
                        for token in llm.stream(prompt):
                            tokens.append(token)
                            slot.markdown(
                                f"<div class='bubble-ai'><div class='lbl'>◆ NarrativeForge</div>"
                                f"<div class='txt'>{''.join(tokens)}</div></div>",
                                unsafe_allow_html=True)
                        story["messages"][i]["content"] = "".join(tokens)
                        save_story(username, story); st.rerun()
                with b2:
                    if st.button("✍️ Continue This", key=f"cont_{i}", use_container_width=True):
                        prompt = _build_prompt(last_content, story, "paragraph", characters, scenes)
                        slot   = st.empty()
                        tokens = []
                        for token in llm.stream(prompt):
                            tokens.append(token)
                            slot.markdown(
                                f"<div class='bubble-ai'><div class='lbl'>◆ NarrativeForge</div>"
                                f"<div class='txt'>{''.join(tokens)}</div></div>",
                                unsafe_allow_html=True)
                        story["messages"].append({"role": "assistant", "content": "".join(tokens)})
                        save_story(username, story); st.rerun()
    else:
        st.markdown("""
            <div style='text-align:center;padding:50px 20px;
                background:linear-gradient(135deg,#F5F6FA,#EDEEF5);
                border:1px solid rgba(77,107,254,0.10);border-radius:20px;margin:20px 0;'>
                <div style='font-size:2rem;margin-bottom:12px;color:#4D6BFE;'>◆</div>
                <div style='color:#A8BCFF;font-size:1.05rem;font-family:'Inter',sans-serif;'>Begin your story below.</div>
                <div style='color:#6B7080;font-size:0.85rem;margin-top:6px;'>Type a line or scene — the AI will stream a continuation.</div>
                <div style='color:#6B7080;font-size:0.8rem;margin-top:4px;'>Add your writing voice in Settings to shape the style.</div>
            </div>
        """, unsafe_allow_html=True)

    # ── Input — disabled while AI is generating ──────────────────────────────
    if not st.session_state.get("_ai_generating"):
        user_input = st.chat_input("Write your story here…")
    else:
        st.chat_input("AI is writing… please wait", disabled=True)
        user_input = None

    if user_input:
        # FIX: Deduplication — ignore if identical to the last user message
        last_msgs = story.get("messages", [])
        last_user_msg = next(
            (m["content"] for m in reversed(last_msgs) if m["role"] == "user"), None
        )
        if user_input.strip() == (last_user_msg or "").strip():
            # Silently ignore the duplicate — don't append, don't rerun
            user_input = None

    if user_input:
        # FIX: Cooldown — ignore if same input sent within last 3 seconds
        last_send_time = st.session_state.get("_last_send_time", 0)
        last_send_text = st.session_state.get("_last_send_text", "")
        now = time.time()
        if user_input.strip() == last_send_text and (now - last_send_time) < 3:
            user_input = None
        else:
            st.session_state["_last_send_time"] = now
            st.session_state["_last_send_text"] = user_input.strip()

    if user_input:
        story["messages"].append({"role": "user", "content": user_input[:4000]})
        save_story(username, story)

        if _is_non_story(user_input):
            story["messages"].append({"role": "assistant", "content": _get_redirect()})
            save_story(username, story)
            st.session_state.show_cowrite_options = False
        else:
            st.session_state.show_cowrite_options = True
            st.session_state.pending_input = user_input
        _bust_cache()
        st.rerun()

    # ── Mode buttons with live streaming ────────────────────────────────────────
    if st.session_state.get("show_cowrite_options"):
        st.markdown("<div class='mode-header'>◆ How should I continue?</div>", unsafe_allow_html=True)
        m1, m2, m3 = st.columns(3)
        m4, m5, m6 = st.columns(3)

        def _run_stream(mode_key):
            # FIX #9: set generating flag so input is disabled during generation
            st.session_state["_ai_generating"] = True
            prompt = _build_prompt(st.session_state.pending_input, story, mode_key, characters, scenes)
            slot   = st.empty()
            tokens = []
            error_occurred = False
            for token in llm.stream(prompt):
                # FIX #11: detect error tokens — don't persist them
                if token.startswith("⚠️") or token.startswith("[Error"):
                    error_occurred = True
                    slot.error(token)
                    break
                tokens.append(token)
                slot.markdown(
                    f"<div class='bubble-ai'><div class='lbl'>◆ NarrativeForge</div>"
                    f"<div class='txt'>{''.join(tokens)}</div></div>",
                    unsafe_allow_html=True)
            # FIX #12: Only save AI response if it has real content (no error messages)
            ai_content = "".join(tokens).strip()
            if ai_content and not error_occurred:
                story["messages"].append({"role": "assistant", "content": ai_content})
                save_story(username, story)
            elif error_occurred and not ai_content:
                # Nothing usable was generated — don't pollute story with blank entry
                pass
            elif ai_content and error_occurred:
                # Partial content before error — save what we have, it's valid prose
                story["messages"].append({"role": "assistant", "content": ai_content})
                save_story(username, story)
            st.session_state["_ai_generating"] = False
            st.session_state.show_cowrite_options = False
            _bust_cache()
            st.rerun()

        # FIX #9: disable all mode buttons while generating
        _gen = st.session_state.get("_ai_generating", False)
        with m1:
            if st.button("Continue", use_container_width=True, type="primary",
                         disabled=_gen):
                _run_stream("continue")
        with m2:
            if st.button("Paragraph", use_container_width=True, type="primary",
                         disabled=_gen):
                _run_stream("paragraph")
        with m3:
            if st.button("Smart", use_container_width=True, type="secondary",
                         disabled=_gen):
                _run_stream("auto")
        with m4:
            if st.button("Dialogue", use_container_width=True, type="primary",
                         disabled=_gen,
                         help="Generate a spoken exchange between characters"):
                _run_stream("dialogue")
        with m5:
            if st.button("Plot Twist", use_container_width=True, type="primary",
                         disabled=_gen,
                         help="Introduce an unexpected revelation"):
                _run_stream("twist")
        with m6:
            if st.button("Monologue", use_container_width=True, type="secondary",
                         disabled=_gen,
                         help="Inner thoughts of the protagonist"):
                _run_stream("monologue")

    # ── Revision Mode toggle ─────────────────────────────────────────────────
    st.markdown("<div style='margin-top:8px;'></div>", unsafe_allow_html=True)
    rev_open = st.session_state.get("show_revision_mode", False)
    if st.button(
        "✂️ Revision Mode" if not rev_open else "✖ Close Revision Mode",
        use_container_width=True,
        help="Paste any paragraph and get 3 AI rewrites in different styles",
        key="toggle_revision"
    ):
        st.session_state["show_revision_mode"] = not rev_open
        st.rerun()

    if st.session_state.get("show_revision_mode"):
        _revision_mode_panel(story, characters)
